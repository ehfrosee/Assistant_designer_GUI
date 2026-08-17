# -*- coding: utf-8 -*-
"""Экспорт в DOCX, XLSX, JSON, CSV"""

import json
import io
import re
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook


class ExportHandlers:
    """Класс-контейнер для методов экспорта в разные форматы."""

    def __init__(self, parent):
        self.parent = parent

    # ---------- Вспомогательные методы ----------
    def _parse_markdown_table(self, table_lines):
        """Парсит строки markdown-таблицы в список списков."""
        lines = [line.strip() for line in table_lines if line.strip()]
        if len(lines) < 2:
            return None
        header_line = lines[0]
        sep_line = lines[1] if len(lines) > 1 else None
        if not sep_line or '---' not in sep_line:
            data_rows = lines
            headers = []
        else:
            headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
            data_rows = lines[2:]
        table = []
        if headers:
            table.append(headers)
        for row_line in data_rows:
            cells = [cell.strip() for cell in row_line.split('|') if cell.strip()]
            if len(cells) < len(headers):
                cells.extend([''] * (len(headers) - len(cells)))
            table.append(cells[:len(headers)])
        return table

    def _convert_csv_to_table(self, text: str) -> list:
        """Преобразует CSV-текст в список списков для таблицы."""
        if not text:
            return None
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if len(lines) < 2:
            return None
        separators = [',', '\t', ';', '|']
        best_sep = None
        best_count = 0
        for sep in separators:
            count = sum(line.count(sep) for line in lines)
            if count > best_count:
                best_count = count
                best_sep = sep
        if best_sep is None or best_count < 2:
            return None
        table = []
        for line in lines:
            if best_sep == ',':
                import csv
                reader = csv.reader(io.StringIO(line))
                row = next(reader)
            else:
                row = [cell.strip() for cell in line.split(best_sep)]
            table.append(row)
        return table

    def _is_likely_csv(self, text: str) -> bool:
        """Определяет, похож ли текст на CSV."""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if len(lines) < 2:
            return False
        separators = [',', '\t', ';', '|']
        for sep in separators:
            count = sum(1 for line in lines if sep in line)
            if count >= len(lines) * 0.5:
                return True
        return False

    def _convert_csv_block_to_table(self, csv_text: str) -> list:
        """Преобразует CSV-текст (с кавычками) в список списков."""
        if not csv_text:
            return None
        import csv
        lines = csv_text.strip().split('\n')
        if len(lines) < 2:
            return None
        try:
            reader = csv.reader(io.StringIO(csv_text))
            table = list(reader)
            if table and len(table) > 0 and len(table[0]) > 0:
                return table
        except:
            pass
        separators = ['\t', ';', '|']
        for sep in separators:
            try:
                table = [line.split(sep) for line in lines]
                if table and len(table) > 0 and len(table[0]) > 1:
                    return table
            except:
                continue
        return None

    def _apply_inline_formatting(self, paragraph, text: str):
        """Применяет inline-форматирование: **жирный**, __жирный__, *курсив*, _курсив_, `код`."""
        import re
        parts = re.split(r'(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`)', text)
        for part in parts:
            if not part:
                continue
            if (part.startswith('**') and part.endswith('**')) or (part.startswith('__') and part.endswith('__')):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                paragraph.add_run(part)

    def _json_to_word_table(self, doc, data):
        """Преобразует JSON-данные в таблицу Word."""
        from docx.shared import Inches
        if isinstance(data, list) and data and all(isinstance(item, dict) for item in data):
            headers = list(data[0].keys())
            table = doc.add_table(rows=1 + len(data), cols=len(headers))
            table.style = 'Table Grid'
            for col_idx, header in enumerate(headers):
                cell = table.cell(0, col_idx)
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for row_idx, item in enumerate(data, 1):
                for col_idx, key in enumerate(headers):
                    val = item.get(key, '')
                    if isinstance(val, (list, dict)):
                        val = json.dumps(val, ensure_ascii=False)
                    cell = table.cell(row_idx, col_idx)
                    cell.text = str(val)
                    cell.width = None
        elif isinstance(data, dict):
            table = doc.add_table(rows=len(data), cols=2)
            table.style = 'Table Grid'
            for row_idx, (key, val) in enumerate(data.items()):
                table.cell(row_idx, 0).text = str(key)
                if isinstance(val, (list, dict)):
                    val = json.dumps(val, ensure_ascii=False)
                table.cell(row_idx, 1).text = str(val)
        else:
            p = doc.add_paragraph()
            run = p.add_run(json.dumps(data, ensure_ascii=False))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

    # ---------- Сохранение чата в DOCX ----------
    def save_chat_as_docx(self, chat_text: str) -> bool:
        """Сохраняет чат в DOCX и возвращает путь или False."""
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from tkinter import filedialog, messagebox

        doc = Document()
        lines = chat_text.split('\n')
        in_code_block = False
        code_block_lines = []
        code_block_lang = ''
        csv_buffer = []
        in_csv = False

        def flush_csv():
            nonlocal csv_buffer, in_csv
            if csv_buffer:
                table = self._convert_csv_to_table('\n'.join(csv_buffer))
                if table and len(table) > 0 and len(table[0]) > 0:
                    word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                    word_table.style = 'Table Grid'
                    for i, row_data in enumerate(table):
                        for j, cell_text in enumerate(row_data):
                            if j < len(row_data):
                                word_table.cell(i, j).text = cell_text
                else:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(csv_buffer))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                csv_buffer = []
                in_csv = False

        def add_paragraph_with_markdown(text: str):
            p = doc.add_paragraph()
            self._apply_inline_formatting(p, text)

        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            i += 1
            if not line:
                if in_csv:
                    csv_buffer.append('')
                else:
                    doc.add_paragraph()
                continue
            if line.startswith('---') and line.endswith('---'):
                continue
            if line.startswith('Вы  ') or line.startswith('Ассистент  ') or line.startswith('*Система*'):
                flush_csv()
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(10)
                continue
            if line.startswith('```'):
                flush_csv()
                if in_code_block:
                    if code_block_lines:
                        block_text = '\n'.join(code_block_lines)
                        if code_block_lang == 'csv':
                            table = self._convert_csv_block_to_table(block_text)
                            if table and len(table) > 0 and len(table[0]) > 0:
                                word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                                word_table.style = 'Table Grid'
                                for r, row_data in enumerate(table):
                                    for c, cell_text in enumerate(row_data):
                                        if c < len(row_data):
                                            word_table.cell(r, c).text = cell_text
                            else:
                                p = doc.add_paragraph()
                                run = p.add_run(block_text)
                                run.font.name = 'Courier New'
                                run.font.size = Pt(9)
                        else:
                            p = doc.add_paragraph()
                            run = p.add_run(block_text)
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                        code_block_lines = []
                        code_block_lang = ''
                    in_code_block = False
                else:
                    in_code_block = True
                    code_block_lang = line[3:].strip().lower()
                continue
            if in_code_block:
                code_block_lines.append(line)
                continue
            if not line.startswith('#') and self._is_likely_csv(line):
                if not in_csv:
                    flush_csv()
                    in_csv = True
                csv_buffer.append(line)
                continue
            elif in_csv:
                flush_csv()
            header_match = re.match(r'^(#+)\s+(.*)$', line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2).strip()
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                if level == 1:
                    run.font.size = Pt(18)
                    p.style = 'Heading 1'
                elif level == 2:
                    run.font.size = Pt(16)
                    p.style = 'Heading 2'
                elif level == 3:
                    run.font.size = Pt(14)
                    p.style = 'Heading 3'
                elif level >= 4:
                    run.font.size = Pt(12)
                    p.style = 'Heading 4'
                continue
            flush_csv()
            add_paragraph_with_markdown(line)

        flush_csv()
        if in_code_block and code_block_lines:
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_block_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Документ Word", "*.docx"), ("Все файлы", "*.*")],
            title="Сохранить чат как DOCX"
        )
        if not file_path:
            return False
        doc.save(file_path)
        return file_path

    # ---------- Сохранение результата в DOCX ----------
    def save_as_docx(self, content: str, fmt: str) -> Document:
        """Преобразует результат в документ DOCX и возвращает Document."""
        doc = Document()
        if fmt == 'markdown':
            self._save_markdown_to_docx(doc, content)
        elif fmt == 'json':
            try:
                if isinstance(content, str):
                    data = json.loads(content)
                else:
                    data = content
                if isinstance(data, list) and data and all(isinstance(item, dict) for item in data):
                    self._json_to_word_table(doc, data)
                elif isinstance(data, dict):
                    self._json_to_word_table(doc, data)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(json.dumps(data, ensure_ascii=False, indent=2))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            except (json.JSONDecodeError, TypeError):
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        else:
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        return doc

    def _save_markdown_to_docx(self, doc, content: str):
        """Сохраняет Markdown-контент в документ DOCX."""
        import re
        lines = content.split('\n')
        in_table = False
        table_lines = []
        in_code_block = False
        code_block_lines = []
        code_block_lang = ''
        csv_buffer = []
        in_csv = False

        def flush_csv():
            nonlocal csv_buffer, in_csv
            if csv_buffer:
                table = self._convert_csv_block_to_table('\n'.join(csv_buffer))
                if table and len(table) > 0 and len(table[0]) > 0:
                    word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                    word_table.style = 'Table Grid'
                    for r, row_data in enumerate(table):
                        for c, cell_text in enumerate(row_data):
                            if c < len(row_data):
                                word_table.cell(r, c).text = cell_text
                else:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(csv_buffer))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                csv_buffer = []
                in_csv = False

        def add_paragraph_with_markdown(text: str):
            p = doc.add_paragraph()
            self._apply_inline_formatting(p, text)

        for line in lines:
            if not line:
                if in_code_block:
                    code_block_lines.append('')
                elif in_csv:
                    csv_buffer.append('')
                else:
                    doc.add_paragraph()
                continue
            if line.startswith('```'):
                flush_csv()
                if in_code_block:
                    if code_block_lines:
                        block_text = '\n'.join(code_block_lines)
                        if code_block_lang == 'csv':
                            table = self._convert_csv_block_to_table(block_text)
                            if table and len(table) > 0 and len(table[0]) > 0:
                                word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                                word_table.style = 'Table Grid'
                                for r, row_data in enumerate(table):
                                    for c, cell_text in enumerate(row_data):
                                        if c < len(row_data):
                                            word_table.cell(r, c).text = cell_text
                                else:
                                    p = doc.add_paragraph()
                                    run = p.add_run(block_text)
                                    run.font.name = 'Courier New'
                                    run.font.size = Pt(9)
                            else:
                                p = doc.add_paragraph()
                                run = p.add_run(block_text)
                                run.font.name = 'Courier New'
                                run.font.size = Pt(9)
                            code_block_lines = []
                            code_block_lang = ''
                        in_code_block = False
                    else:
                        in_code_block = True
                        code_block_lang = line[3:].strip().lower()
                    continue
            if in_code_block:
                code_block_lines.append(line)
                continue
            if not line.startswith('#') and '|' not in line and self._is_likely_csv(line):
                if not in_csv:
                    flush_csv()
                    in_csv = True
                csv_buffer.append(line)
                continue
            elif in_csv:
                flush_csv()
            if '|' in line and not line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                table_lines.append(line)
            else:
                if in_table:
                    if table_lines:
                        table = self._parse_markdown_table(table_lines)
                        if table and len(table) > 0 and len(table[0]) > 0:
                            word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                            word_table.style = 'Table Grid'
                            for r, row_data in enumerate(table):
                                for c, cell_text in enumerate(row_data):
                                    if c < len(row_data):
                                        word_table.cell(r, c).text = cell_text
                    table_lines = []
                    in_table = False
                header_match = re.match(r'^(#+)\s+(.*)$', line)
                if header_match:
                    level = len(header_match.group(1))
                    text = header_match.group(2).strip()
                    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.bold = True
                    if level == 1:
                        run.font.size = Pt(18)
                        p.style = 'Heading 1'
                    elif level == 2:
                        run.font.size = Pt(16)
                        p.style = 'Heading 2'
                    elif level == 3:
                        run.font.size = Pt(14)
                        p.style = 'Heading 3'
                    elif level >= 4:
                        run.font.size = Pt(12)
                        p.style = 'Heading 4'
                    continue
                add_paragraph_with_markdown(line)

        flush_csv()
        if in_table and table_lines:
            table = self._parse_markdown_table(table_lines)
            if table and len(table) > 0 and len(table[0]) > 0:
                word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                word_table.style = 'Table Grid'
                for r, row_data in enumerate(table):
                    for c, cell_text in enumerate(row_data):
                        if c < len(row_data):
                            word_table.cell(r, c).text = cell_text

    # ---------- Сохранение в XLSX ----------
    def save_as_xlsx(self, content: str, fmt: str) -> Workbook:
        """Преобразует результат в XLSX и возвращает Workbook."""
        wb = Workbook()
        if fmt in ('json', 'csv'):
            try:
                if fmt == 'json':
                    data = json.loads(content)
                    if isinstance(data, dict) and 'systems' in data:
                        data = data['systems']
                    if isinstance(data, list):
                        df = pd.DataFrame(data)
                    else:
                        df = pd.DataFrame([data])
                else:
                    df = pd.read_csv(io.StringIO(content))
                with pd.ExcelWriter('temp.xlsx', engine='openpyxl') as writer:
                    df.to_excel(writer, index=False, sheet_name='Data')
                wb = Workbook()
                for sheet in writer.sheets:
                    for row in writer.sheets[sheet].iter_rows():
                        wb.active.append([cell.value for cell in row])
                return wb
            except Exception:
                pass
        if fmt == 'markdown':
            lines = content.split('\n')
            table_lines = []
            in_table = False
            text_parts = []
            tables = []
            for line in lines:
                if '|' in line and not line.strip().startswith('|'):
                    if not in_table:
                        in_table = True
                    table_lines.append(line)
                else:
                    if in_table:
                        if table_lines:
                            table = self._parse_markdown_table(table_lines)
                            if table:
                                tables.append(table)
                        table_lines = []
                        in_table = False
                        text_parts.append(line)
                    else:
                        text_parts.append(line)
            if table_lines:
                table = self._parse_markdown_table(table_lines)
                if table:
                    tables.append(table)
            if text_parts:
                ws_text = wb.active
                ws_text.title = "Text"
                ws_text.cell(row=1, column=1, value='\n'.join(text_parts))
            if tables:
                ws_table = wb.create_sheet("Tables")
                start_row = 1
                for table in tables:
                    for row_data in table:
                        for col_idx, cell_value in enumerate(row_data, 1):
                            ws_table.cell(row=start_row, column=col_idx, value=cell_value)
                        start_row += 1
                    start_row += 1
        else:
            ws = wb.active
            ws.cell(row=1, column=1, value=content)
        return wb