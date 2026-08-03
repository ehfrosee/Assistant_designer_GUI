# -*- coding: utf-8 -*-
"""Главное окно приложения AI Document Analyst (Chat)"""

import logging
import threading
import os
import glob
import json
import io
import re
from pathlib import Path
from datetime import datetime
from typing import Optional

import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, Frame, Label, Text, Toplevel, ttk, Menu
from tkinter import END, DISABLED, NORMAL, LEFT, RIGHT, X, Y, BOTH
import ttkbootstrap as tb
from ttkbootstrap.constants import *
import markdown
import pandas as pd
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ---- Импорты из других модулей проекта ----
from config.config_manager import ConfigManager
from core.chat_client import ChatGPTClient
from gui.file_loader import FileLoader
from gui.dialog_manager import DialogManager
from orchestrator.dispatcher import run_analysis
from gui.scenario_editor import launch_scenario_editor

# Импорты вынесенных виджетов
from gui.chat_widget import ChatWidget
from gui.settings_dialog import SettingsDialog


class MainWindow(tb.Window):
    """Главное окно приложения, объединяющее чат, анализ документов по сценариям, команды и настройки"""

    # ---------- Инициализация и настройка ----------
    def __init__(self, config_manager: ConfigManager):
        theme = config_manager.config.get("ui", {}).get("theme", "light")
        theme_name = "darkly" if theme == "dark" else "cosmo"
        super().__init__(themename=theme_name)

        self.config_manager = config_manager
        self.chat_client = ChatGPTClient(config_manager)
        self.dialog_manager = DialogManager(config_manager, self.chat_client)
        self._stream_buffer = ""
        self.logger = logging.getLogger("main_window")
        self.analysis_thread = None
        self.cancel_event = None
        self.result_format = "text"
        self.setup_logging()
        self.init_ui()
        self.bind_signals()
        self.check_api_key()
        self.after(500, self.test_api_connection)
        self.after(1000, self._auto_load_prompt)

        self.available_models = []
        self.load_models_at_startup()
        self.attached_files = []  # список кортежей (filename, content)
        self.doc_paths_var = []  # список файлов для анализа (несколько)

    def setup_logging(self):
        """Настройка логирования в файл и консоль"""
        logs_dir = self.config_manager.config.get("paths", {}).get("logs_dir", "./logs")
        Path(logs_dir).mkdir(exist_ok=True)
        log_file = Path(logs_dir) / "app.log"
        logging.basicConfig(
            level=logging.INFO,
            format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
            handlers=[
                logging.FileHandler(log_file, encoding="utf-8"),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger("main_window")
        logging.getLogger("chat_client").setLevel(logging.DEBUG)

    def _safe_filename(self, title: str, default: str = "untitled") -> str:
        """Преобразует строку в безопасное имя файла (убирает недопустимые символы)."""
        import re
        # Убираем символы, запрещённые в именах файлов Windows/Linux
        safe = re.sub(r'[<>:"/\\|?*]', '_', title.strip())
        # Убираем лишние пробелы и точки в конце
        safe = safe.strip(' .')
        if not safe:
            safe = default
        # Ограничиваем длину
        if len(safe) > 100:
            safe = safe[:100]
        return safe

    def _save_selected_as_file(self):
        """Сохраняет выделенный текст чата в один файл с именем из первой строки."""
        selected = self.chat_widget.get_selected_text()
        if not selected:
            messagebox.showwarning("Сохранение", "Ничего не выделено в чате.")
            return

        # Первая строка (до первого переноса) как имя файла
        first_line = selected.splitlines()[0].strip() if selected else ""
        if not first_line:
            first_line = "selected_text"

        base_name = self._safe_filename(first_line, "selected_text")
        default_ext = ".md" if first_line.startswith('#') else ".txt"

        file_path = filedialog.asksaveasfilename(
            defaultextension=default_ext,
            filetypes=[("Текстовые файлы", "*.txt"), ("Markdown", "*.md"), ("Все файлы", "*.*")],
            initialfile=base_name + default_ext,
            title="Сохранить выделенный текст"
        )
        if not file_path:
            return

        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(selected)
            self.statusbar.config(text=f"Выделенный текст сохранён в {file_path}")
            self.add_system_message(f"📄 Выделенный текст сохранён в {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка сохранения", f"Не удалось сохранить файл:\n{e}")

    def _save_selected_as_sections(self):
        """
        Сохраняет выделенный текст, разделяя по строкам, начинающимся с "---" и содержащим имя файла.
        Строки, состоящие только из "---" (без имени), игнорируются как разделители.
        Каждый раздел сохраняется в отдельный файл.
        Имя файла берётся из строки "--- <имя_файла>".
        Удаляет обратные кавычки (```) и возможный язык в начале и конце содержимого.
        """
        selected = self.chat_widget.get_selected_text()
        if not selected:
            messagebox.showwarning("Сохранение", "Ничего не выделено в чате.")
            return

        import re

        # Ищем строки, начинающиеся с "---" и содержащие имя файла (не только пробелы)
        pattern = r'(?=^[ \t]*---\s+\S+.*$)'
        sections = re.split(pattern, selected, flags=re.MULTILINE)
        sections = [s.strip() for s in sections if s.strip()]

        if not sections:
            messagebox.showwarning("Сохранение",
                                   "Выделенный текст не содержит маркеров разделов (строк, начинающихся с '---' и имени файла).")
            return

        if len(sections) == 1:
            reply = messagebox.askyesno("Сохранение",
                                        "Текст не разбит на разделы (только один маркер).\n"
                                        "Сохранить как один файл?")
            if reply:
                self._save_selected_as_file()
            return

        folder_path = filedialog.askdirectory(title="Выберите папку для сохранения разделов")
        if not folder_path:
            return

        saved_count = 0
        errors = []

        for section in sections:
            lines = section.splitlines()
            if not lines:
                continue

            # Первая строка должна начинаться с "---" и содержать имя файла
            first_line = lines[0].strip()
            if not first_line.startswith('---'):
                # Если маркер отсутствует (например, первый кусок текста до первого маркера), пропускаем
                # или можно сохранить как unnamed, но лучше пропустить, так как это не полная секция
                continue

            # Извлекаем имя файла после "---"
            file_name_part = first_line[3:].strip()
            file_name_part = file_name_part.strip('"').strip("'")
            if not file_name_part:
                # Если имя пустое (что маловероятно при новом pattern), пропускаем
                continue

            title = file_name_part
            # Остальные строки (кроме первой) — содержимое
            content = "\n".join(lines[1:]).strip()

            # Удаляем обратные кавычки и возможный язык в начале/конце
            if content.startswith('```') and content.endswith('```'):
                lines_content = content.splitlines()
                if lines_content and lines_content[0].strip().startswith('```'):
                    lines_content = lines_content[1:]
                if lines_content and lines_content[-1].strip() == '```':
                    lines_content = lines_content[:-1]
                content = "\n".join(lines_content).strip()

            # Дополнительная очистка от обратных кавычек
            content = re.sub(r'^```[a-zA-Z]*\s*', '', content)
            content = re.sub(r'\s*```$', '', content)

            # Формируем имя файла
            base_name = self._safe_filename(title, "section")
            # Если имя не содержит расширения, добавляем .txt
            if '.' not in os.path.splitext(base_name)[1]:
                base_name += ".txt"

            file_path = os.path.join(folder_path, base_name)

            counter = 1
            while os.path.exists(file_path):
                name, ext = os.path.splitext(base_name)
                file_path = os.path.join(folder_path, f"{name}_{counter}{ext}")
                counter += 1

            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                saved_count += 1
            except Exception as e:
                errors.append(f"{file_path}: {e}")

        if saved_count > 0:
            self.statusbar.config(text=f"Сохранено {saved_count} разделов в {folder_path}")
            self.add_system_message(f"📁 Сохранено {saved_count} разделов в {folder_path}")
        if errors:
            self.add_system_message(f"⚠️ Ошибки при сохранении:\n" + "\n".join(errors))

    # ---------- Построение пользовательского интерфейса ----------
    def init_ui(self):
        """Создаёт все элементы главного окна: меню, вкладки, панели"""
        self.title("AI Document Analyst (Chat)")
        self.geometry("900x600")
        self.minsize(600, 400)

        # ---- Меню ----
        menubar = tb.Menu(self)
        self.config(menu=menubar)

        # Меню "Файл"
        file_menu = tb.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Файл", menu=file_menu)
        file_menu.add_command(label="Выход", command=self.destroy)

        # Меню "Диалог"
        dialog_menu = tb.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Диалог", menu=dialog_menu)
        dialog_menu.add_command(label="Новый диалог", command=self.new_dialog)
        dialog_menu.add_command(label="Сохранить диалог", command=self.save_dialog)
        dialog_menu.add_command(label="Загрузить диалог", command=self.load_dialog)
        dialog_menu.add_command(label="Обобщить диалог (summary)", command=self.summarize_dialog)
        dialog_menu.add_separator()
        dialog_menu.add_command(label="Очистить загруженные документы", command=self._cmd_clear_documents)
        dialog_menu.add_command(label="Очистить загруженные сценарии", command=self._cmd_clear_scenarios)
        dialog_menu.add_separator()
        dialog_menu.add_command(label="Загрузить промпт", command=self._load_prompt_global)
        dialog_menu.add_command(label="Сохранить промпт", command=self._save_prompt_global)

        # Меню "Инструменты"
        tools_menu = tb.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Инструменты", menu=tools_menu)
        tools_menu.add_command(label="Копировать чат", command=self.copy_chat)
        tools_menu.add_command(label="Сохранить чат в DOCX", command=self._save_chat_as_docx_wrapper)
        tools_menu.add_separator()
        tools_menu.add_command(label="Сохранить выделенное как файл", command=self._save_selected_as_file)
        tools_menu.add_command(label="Сохранить выделенное по разделам", command=self._save_selected_as_sections)
        tools_menu.add_separator()
        tools_menu.add_command(label="Очистить чат", command=self.clear_chat)
        tools_menu.add_command(label="Настройки", command=self.open_settings)

        # Меню "Команды"
        commands_menu = tb.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Команды", menu=commands_menu)
        commands_menu.add_command(label="📄 Загрузить документ", command=self._cmd_load_document_dialog)
        commands_menu.add_command(label="📋 Загрузить сценарий", command=self._cmd_load_scenario_dialog)
        commands_menu.add_separator()
        commands_menu.add_command(label="▶️ Запустить сценарий", command=self._cmd_run_scenario)
        commands_menu.add_command(label="⏹️ Остановить сценарий", command=self._cmd_stop_scenario)
        commands_menu.add_separator()
        commands_menu.add_command(label="📁 Импортировать папку", command=self._cmd_import_folder_dialog)

        # Подменю "Экспортировать результат"
        export_menu = tb.Menu(commands_menu, tearoff=0)
        commands_menu.add_cascade(label="💾 Экспортировать результат", menu=export_menu)
        for fmt, label in [("json", "JSON"), ("csv", "CSV"), ("md", "Markdown (MD)"), ("docx", "Word (DOCX)"), ("xlsx", "Excel (XLSX)")]:
            export_menu.add_command(label=label, command=lambda f=fmt: self._cmd_export_result(f))

        commands_menu.add_separator()
        commands_menu.add_command(label="🗑️ Очистить документы", command=self._cmd_clear_documents)
        commands_menu.add_command(label="🗑️ Очистить сценарии", command=self._cmd_clear_scenarios)
        commands_menu.add_separator()
        commands_menu.add_command(label="❓ Помощь", command=self._cmd_help)

        # Меню "Помощь"
        help_menu = tb.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Помощь", menu=help_menu)
        help_menu.add_command(label="О программе", command=self.show_about)

        # ---- Основной контент: вкладки ----
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill=BOTH, expand=True)

        # Вкладка "Чат"
        self.chat_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.chat_frame, text="Чат")
        self.chat_widget = ChatWidget(self.chat_frame, self.on_user_message, self)
        self.chat_widget.pack(fill=BOTH, expand=True)

        # Вкладка "Анализ документа"
        self.analysis_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.analysis_frame, text="Анализ документа")
        self._build_analysis_tab()

        # ---- Панель загрузки файлов и кнопок чата ----
        file_frame = Frame(self)
        file_frame.pack(fill=X, padx=5, pady=5)

        self.attach_btn = tb.Button(file_frame, text="📎 Прикрепить файл", bootstyle="secondary", command=self.load_file)
        self.attach_btn.pack(side=LEFT, padx=2)

        self.file_label = tb.Label(file_frame, text="Файл не загружен", bootstyle="info")
        self.file_label.pack(side=LEFT, padx=10)

        self.clear_file_btn = tb.Button(file_frame, text="Убрать файл", bootstyle="danger", command=self.clear_file, state=DISABLED)
        self.clear_file_btn.pack(side=LEFT, padx=2)

        self.copy_chat_btn = tb.Button(file_frame, text="📋 Копировать чат", bootstyle="info", command=self.copy_chat)
        self.copy_chat_btn.pack(side=RIGHT, padx=2)

        self.save_chat_docx_btn = tb.Button(file_frame, text="💾 Сохранить чат (DOCX)", bootstyle="info", command=self._save_chat_as_docx_wrapper)
        self.save_chat_docx_btn.pack(side=RIGHT, padx=2)

        # Статусная строка
        self.statusbar = tb.Label(self, text="Готов", bootstyle="info", anchor=W)
        self.statusbar.pack(side=BOTTOM, fill=X)

        # Приветственное сообщение
        self.chat_widget.add_message("Приложение запущено. Введите вопрос.", "assistant")

    def _build_analysis_tab(self):
        """Строит интерфейс вкладки 'Анализ документа'"""
        # Панель выбора документа
        doc_frame = ttk.LabelFrame(self.analysis_frame, text="Документ", padding=5)
        doc_frame.pack(fill=X, padx=5, pady=5)
        self.doc_path_var = tb.StringVar()
        self.doc_entry = tb.Entry(doc_frame, textvariable=self.doc_path_var, state="readonly")
        self.doc_entry.pack(side=LEFT, fill=X, expand=True, padx=(0, 5))
        tb.Button(doc_frame, text="Выбрать файл", bootstyle="secondary", command=self._select_document).pack(side=RIGHT)

        # Панель выбора сценария
        scenario_frame = ttk.LabelFrame(self.analysis_frame, text="Сценарий", padding=5)
        scenario_frame.pack(fill=X, padx=5, pady=5)
        self.scenario_path_var = tb.StringVar()
        self.scenario_entry = tb.Entry(scenario_frame, textvariable=self.scenario_path_var, state="readonly")
        self.scenario_entry.pack(side=LEFT, fill=X, expand=True, padx=(0, 5))
        tb.Button(scenario_frame, text="Загрузить JSON", bootstyle="secondary", command=self._select_scenario).pack(side=RIGHT, padx=2)
        tb.Button(scenario_frame, text="Редактор сценариев", bootstyle="info", command=self._open_scenario_editor).pack(side=RIGHT, padx=2)

        # Кнопки управления
        control_frame = ttk.Frame(self.analysis_frame)
        control_frame.pack(fill=X, padx=5, pady=5)
        self.analyze_btn = tb.Button(control_frame, text="Запустить анализ", bootstyle="success", command=self._start_analysis, state=DISABLED)
        self.analyze_btn.pack(side=LEFT, padx=2)
        self.cancel_btn = tb.Button(control_frame, text="Отменить", bootstyle="danger", command=self._cancel_analysis, state=DISABLED)
        self.cancel_btn.pack(side=LEFT, padx=2)

        # Прогресс
        self.progress_var = tb.IntVar(value=0)
        self.progress_bar = ttk.Progressbar(self.analysis_frame, variable=self.progress_var, mode='determinate')
        self.progress_bar.pack(fill=X, padx=5, pady=5)

        # Вкладки вывода
        output_notebook = ttk.Notebook(self.analysis_frame)
        output_notebook.pack(fill=BOTH, expand=True, padx=5, pady=5)

        # Лог выполнения
        log_frame = ttk.Frame(output_notebook)
        output_notebook.add(log_frame, text="Лог выполнения")
        self.log_text = scrolledtext.ScrolledText(log_frame, height=8, wrap=WORD, state=DISABLED)
        self.log_text.pack(fill=BOTH, expand=True)

        # Результат анализа
        result_frame = ttk.Frame(output_notebook)
        output_notebook.add(result_frame, text="Результат анализа")
        self.result_text = scrolledtext.ScrolledText(result_frame, height=10, wrap=WORD, state=DISABLED)
        self.result_text.pack(fill=BOTH, expand=True)
        btn_result_frame = ttk.Frame(result_frame)
        btn_result_frame.pack(fill=X, pady=5)
        tb.Button(btn_result_frame, text="Копировать результат", bootstyle="secondary", command=self._copy_result).pack(side=LEFT, padx=2)
        tb.Button(btn_result_frame, text="Сохранить результат", bootstyle="secondary", command=self._save_result).pack(side=LEFT, padx=2)

        # Промежуточные результаты
        intermediate_frame = ttk.Frame(output_notebook)
        output_notebook.add(intermediate_frame, text="Промежуточные результаты")
        self.intermediate_combo = ttk.Combobox(intermediate_frame, state="readonly", width=30)
        self.intermediate_combo.pack(fill=X, pady=(5, 5))
        self.intermediate_combo.bind("<<ComboboxSelected>>", self._on_intermediate_selected)
        self.intermediate_text = scrolledtext.ScrolledText(intermediate_frame, height=8, wrap=WORD, state=DISABLED)
        self.intermediate_text.pack(fill=BOTH, expand=True)

        self.analysis_thread = None
        self.cancel_event = None

    # ---------- Сигналы и колбэки ----------
    def bind_signals(self):
        """Подключает колбэки для ChatGPTClient"""
        self.chat_client.on_response = self._on_ai_response
        self.chat_client.on_chunk = self._on_ai_chunk
        self.chat_client.on_error = self._on_api_error
        self.chat_client.on_start = self._on_request_start
        self.chat_client.on_finish = self._on_request_finish

    # ---------- Обработчики событий чата ----------
    def on_user_message(self, text: str):
        """
        Основной обработчик сообщений пользователя:
        - если команда (/...) – передаёт в _handle_command
        - иначе формирует запрос к ChatGPT с учётом контекста и загруженных файлов
        """
        if text.startswith('/'):
            self._handle_command(text)
            return

        self.chat_widget.add_message(text, "user")
        self.dialog_manager.add_message("user", text)

        # Сбор контекста
        dialog_params = self.config_manager.get_dialog_params()
        history_pairs = dialog_params.get("history_pairs", 3)
        pairs = self.dialog_manager.get_conversation_pairs(history_pairs)
        history_text = ""
        for i, (q, a) in enumerate(pairs):
            history_text += f"Вопрос {i + 1}: {q}\nОтвет {i + 1}: {a}\n"
        summary = self.dialog_manager.current_dialog.get("summary", "")

        # Содержимое загруженных файлов
        files_text = ""
        loaded_docs = self.dialog_manager.get_loaded_documents()
        for doc in loaded_docs:
            files_text += f"--- {doc['filename']} ---\n{doc['content']}\n"
        for fname, content in self.attached_files:
            if not any(doc["filename"] == fname for doc in loaded_docs):
                files_text += f"--- {fname} ---\n{content}\n"

        # Загрузка промптов
        prompts = self.chat_client._load_prompts()
        system_prompt = prompts.get("dialog_system_prompt", "Ты полезный ассистент.")
        user_prompt_template = prompts.get(
            "dialog_user_prompt",
            "Краткое содержание предыдущих сообщений (summary):\n{summary}\n\n"
            "Последние вопросы и ответы:\n{history}\n\n"
            "Содержимое загруженных файлов:\n{files}\n\n"
            "Текущий вопрос пользователя: {question}"
        )
        user_prompt = user_prompt_template.format(
            summary=summary,
            history=history_text,
            files=files_text,
            question=text
        )

        # Проверка размера файла (рекомендация, если слишком большой)
        if files_text:
            import tiktoken
            try:
                model = self.config_manager.get_llm_params().get("model", "gpt-4o-mini")
                encoding = tiktoken.encoding_for_model(model)
                total_tokens = len(encoding.encode(files_text))
                self.logger.info(f"📊 Размер файла в токенах: {total_tokens}")
                max_safe_tokens = 100000 if ("gpt-4" in model or "gpt-5" in model) else 50000
                if total_tokens > max_safe_tokens:
                    recommendation = self._recommend_file_split(files_text, model)
                    if recommendation:
                        self.chat_widget.add_message(recommendation, "assistant")
                        self.dialog_manager.add_message("assistant", recommendation)
                    else:
                        self.chat_widget.add_message("⚠️ Файл слишком большой. Попробуйте разбить его на части и загружать по отдельности.", "assistant")
                    return
            except Exception as e:
                self.logger.warning(f"Ошибка при оценке размера файла: {e}")

        # Отправка запроса
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        self._stream_buffer = ""
        self.chat_client.send_message(messages, stream=True)

    # ---------- Обработка команд (/...) ----------
    def _handle_command(self, text: str):
        parts = text.split(maxsplit=1)
        cmd = parts[0].lower()
        args = parts[1] if len(parts) > 1 else ""

        commands = {
            "/help": self._cmd_help,
            "/load_document": lambda: self._cmd_load_document(args),
            "/load_scenario": lambda: self._cmd_load_scenario(args),
            "/run_scenario": self._cmd_run_scenario,
            "/stop_scenario": self._cmd_stop_scenario,  # <-- добавить
            "/export_result": lambda: self._cmd_export_result(args),
            "/import_folder": lambda: self._cmd_import_folder(args),
            "/clear_documents": self._cmd_clear_documents,
            "/clear_scenarios": self._cmd_clear_scenarios,
        }
        if cmd in commands:
            commands[cmd]()
        else:
            self.chat_widget.add_message(f"❌ Неизвестная команда: {cmd}. Введите /help.", "assistant")
            self.dialog_manager.add_message("assistant", f"Неизвестная команда: {cmd}")

    # ---------- Реализация команд (меню и чат) ----------
    def _cmd_help(self):
        help_text = (
            "**Доступные команды:**\n\n"
            "`/load_document <путь>` — загрузить документ в контекст чата\n"
            "`/load_scenario <путь>` — загрузить JSON-сценарий\n"
            "`/run_scenario` — выполнить загруженный сценарий на загруженном документе\n"
            "`/export_result <формат>` — экспортировать результат (json, csv, md, docx, xlsx)\n"
            "`/import_folder <путь>` — загрузить все документы из папки\n"
            "`/clear_documents` — очистить загруженные документы\n"
            "`/clear_scenarios` — очистить загруженные сценарии\n"
            "`/help` — показать эту справку"
        )
        self.chat_widget.add_message(help_text, "assistant")
        self.dialog_manager.add_message("assistant", help_text)

    def _cmd_load_document(self, file_path: str):
        """Загружает один или несколько документов (поддерживает пути с пробелами и разделители)"""
        if not file_path:
            self.chat_widget.add_message("❌ Укажите путь к файлу: /load_document <путь>", "assistant")
            return
        file_path = file_path.strip()
        if os.path.isdir(file_path):
            self._cmd_import_folder(file_path)
            return
        # Множественные пути через запятую или точку с запятой
        if ',' in file_path or ';' in file_path:
            separator = ',' if ',' in file_path else ';'
            paths = [p.strip().strip('"').strip("'") for p in file_path.split(separator) if p.strip()]
            count = 0
            for p in paths:
                p = p.strip('"').strip("'")
                if os.path.exists(p):
                    self._load_single_document(p)
                    count += 1
                else:
                    self.chat_widget.add_message(f"⚠️ Файл не найден: {p}", "assistant")
            if count > 0:
                self.chat_widget.add_message(f"✅ Загружено {count} документов", "assistant")
            return
        # Одиночный файл
        file_path = file_path.strip('"').strip("'")
        self._load_single_document(file_path)

    def _load_single_document(self, file_path: str):
        if not os.path.exists(file_path):
            self.chat_widget.add_message(f"⚠️ Файл не найден: {file_path}", "assistant")
            return
        try:
            result = FileLoader._read_file(file_path, self.config_manager)
            if result:
                filename, content, file_format = result
                self.dialog_manager.add_loaded_document(filename, content, file_format, file_path)
                self.attached_files.append((filename, content))
                self._update_file_label()
                self.chat_widget.add_message(f"✅ Документ '{filename}' загружен в контекст", "assistant")
                self.dialog_manager.add_message("assistant", f"Документ '{filename}' загружен в контекст")
                # Сохраняем конвертированный файл в папку сценария, если сценарий загружен
                scenario_path = getattr(self, 'scenario_path_var', None)
                if scenario_path and scenario_path.get():
                    scenario_dir = Path(scenario_path.get()).parent
                    input_dir = scenario_dir / "input"
                    input_dir.mkdir(parents=True, exist_ok=True)
                    md_filename = Path(filename).stem + ".md"
                    md_path = input_dir / md_filename
                    with open(md_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    self.chat_widget.add_message(f"*Система*: Конвертированный файл сохранён в {md_path}", "assistant")
            else:
                self.chat_widget.add_message(f"❌ Не удалось загрузить документ: {file_path}", "assistant")
        except Exception as e:
            self.chat_widget.add_message(f"❌ Ошибка загрузки: {e}", "assistant")

    def _cmd_load_scenario(self, file_path: str):
        if not file_path:
            self.chat_widget.add_message("❌ Укажите путь к сценарию: /load_scenario <путь>", "assistant")
            return
        file_path = file_path.strip().strip('"').strip("'")
        if not os.path.exists(file_path):
            self.chat_widget.add_message(f"❌ Файл сценария не найден: {file_path}", "assistant")
            return
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if 'stages' not in data:
                    raise ValueError("Нет ключа 'stages'")
            self.dialog_manager.add_loaded_scenario(file_path, data)
            self.chat_widget.add_message(f"✅ Сценарий '{Path(file_path).stem}' загружен", "assistant")
            self.dialog_manager.add_message("assistant", f"Сценарий '{Path(file_path).stem}' загружен")
        except Exception as e:
            self.chat_widget.add_message(f"❌ Ошибка загрузки сценария: {e}", "assistant")

    def _cmd_run_scenario(self):
        scenario = self.dialog_manager.get_last_scenario()
        document = self.dialog_manager.get_last_document()
        if not scenario:
            self.chat_widget.add_message("❌ Сценарий не загружен. Используйте /load_scenario", "assistant")
            return
        if not document:
            self.chat_widget.add_message("❌ Документ не загружен. Используйте /load_document", "assistant")
            return
        doc_path = document.get("path")
        doc_paths = [doc_path]
        scenario_path = scenario.get("path")
        if not doc_path or not os.path.exists(doc_path):
            self.chat_widget.add_message("⚠️ Путь к документу не найден. Используйте /load_document с полным путём", "assistant")
            return
        if not scenario_path or not os.path.exists(scenario_path):
            self.chat_widget.add_message("⚠️ Путь к сценарию не найден. Используйте /load_scenario с полным путём", "assistant")
            return
        api_key = self.config_manager.get_api_key()
        if not api_key:
            self.chat_widget.add_message("❌ API ключ не найден. Задайте его в настройках.", "assistant")
            return
        llm_params = self.config_manager.get_llm_params()
        api_settings = {
            "api_key": api_key,
            "model": llm_params.get("model", "gpt-4o-mini"),
            "temperature": llm_params.get("temperature", 0.2),
            "max_tokens": llm_params.get("max_tokens", 2000),
            "timeout": llm_params.get("timeout", 60)
        }
        self.chat_widget.add_message("🔄 Запуск анализа по сценарию...", "assistant")
        self.dialog_manager.add_message("assistant", "Запуск анализа по сценарию...")
        self.cancel_event = threading.Event()
        self.analysis_thread = threading.Thread(
            target=self._run_scenario_thread,
            args=(doc_paths, scenario_path, api_settings),
            daemon=True
        )
        self.analysis_thread.start()

    def _run_scenario_thread(self, doc_paths, scenario_path, api_settings):
        def chat_callback(msg_type, data):
            if msg_type == "log":
                self.after(0, lambda: self._append_chat_message(f"*Система*: {data}"))
            elif msg_type == "stage":
                current = data.get("current", 0)
                total = data.get("total", 1)
                name = data.get("name", "unknown")
                self.after(0, lambda: self._append_chat_message(f"*Система*: Этап {current}/{total}: {name}"))
            elif msg_type == "stage_result":
                stage_id = data.get("stage_id", "unknown")
                formatted = data.get("formatted_content")
                if formatted:
                    self.after(0, lambda: self._append_chat_message(
                        f"*Система*: Результат этапа '{stage_id}':\n{formatted}"))
                else:
                    full_result = data.get("full_result", "")
                    self.after(0, lambda: self._append_chat_message(
                        f"*Система*: Результат этапа '{stage_id}':\n{full_result}"))

        try:
            result = run_analysis(
                document_paths=doc_paths,
                scenario_path=scenario_path,
                api_settings=api_settings,
                progress_callback=chat_callback,
                cancel_event=self.cancel_event
            )
            # result должен быть словарём с ключом 'status'
            self.after(0, lambda: self._on_scenario_finished(result))
        except Exception as e:
            self.after(0, lambda: self._append_chat_message(f"❌ Критическая ошибка: {e}"))

    def _append_chat_message(self, text: str):
        self.chat_widget.add_message(text, "assistant")
        self.dialog_manager.add_message("assistant", text)

    def _on_scenario_finished(self, result):
        """Обрабатывает завершение выполнения сценария."""
        # Защита от пустого результата
        if result is None:
            self._append_chat_message("❌ Ошибка: результат анализа пустой")
            self.analysis_thread = None
            return

        # Защита от отсутствия ключа 'status'
        if not isinstance(result, dict):
            self._append_chat_message(f"❌ Ошибка: неверный формат результата: {result}")
            self.analysis_thread = None
            return

        status = result.get("status")

        if status == "success":
            self.dialog_manager.set_last_result(result["result"], result.get("format", "text"))
            self._append_chat_message(f"✅ Анализ завершён успешно! Результат в формате {result.get('format', 'text')}")
            # Показываем результат в чате (если не слишком большой)
            result_text = result.get("result", "")
            self._append_chat_message(f"**Результат:**\n{result_text}")
        elif status == "cancelled":
            self._append_chat_message("⚠️ Анализ отменён пользователем")
        else:
            error_msg = result.get("message", "Неизвестная ошибка")
            self._append_chat_message(f"❌ Ошибка: {error_msg}")

        self.analysis_thread = None

    def add_system_message(self, text: str):
        """
        Добавляет системное сообщение в чат без отправки в GPT.
        Используется для уведомлений о загрузке/сохранении промптов и других системных событиях.
        """
        self.chat_widget.add_message(f"*Система*: {text}", "assistant")
        self.dialog_manager.add_message("assistant", f"*Система*: {text}")

    def _cmd_export_result(self, format_type: str):
        if not format_type:
            self.chat_widget.add_message("❌ Укажите формат: /export_result <json|csv|md|docx|xlsx>", "assistant")
            return
        format_type = format_type.strip().lower()
        valid = {"json", "csv", "md", "docx", "xlsx"}
        if format_type not in valid:
            self.chat_widget.add_message(f"❌ Неподдерживаемый формат: {format_type}. Доступные: {', '.join(valid)}",
                                         "assistant")
            return
        result = self.dialog_manager.get_last_result()
        if not result:
            self.chat_widget.add_message("❌ Нет сохранённых результатов анализа", "assistant")
            return

        ext_map = {"json": ".json", "csv": ".csv", "md": ".md", "docx": ".docx", "xlsx": ".xlsx"}
        file_path = filedialog.asksaveasfilename(
            defaultextension=ext_map[format_type],
            filetypes=[(f"{format_type.upper()} files", f"*{ext_map[format_type]}"), ("All files", "*.*")],
            title=f"Сохранить результат как {format_type.upper()}"
        )
        if not file_path:
            return

        try:
            ext = Path(file_path).suffix.lower()
            content = result["result"]
            if ext == '.docx':
                self._save_as_docx(file_path, content, result.get("format", "text"))
            elif ext == '.xlsx':
                self._save_as_xlsx(file_path, content, result.get("format", "text"))
            elif ext == '.json':
                try:
                    data = json.loads(content)
                    with open(file_path, 'w', encoding='utf-8') as f:
                        json.dump(data, f, ensure_ascii=False, indent=2)
                except json.JSONDecodeError:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                self.chat_widget.add_message(f"✅ Результат сохранён: {file_path}", "assistant")
                self.dialog_manager.add_message("assistant", f"Результат сохранён: {file_path}")
                self.statusbar.config(text=f"Результат сохранён: {file_path}")
            else:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.chat_widget.add_message(f"✅ Результат сохранён: {file_path}", "assistant")
                self.dialog_manager.add_message("assistant", f"Результат сохранён: {file_path}")
                self.statusbar.config(text=f"Результат сохранён: {file_path}")
        except Exception as e:
            self.chat_widget.add_message(f"❌ Ошибка сохранения: {e}", "assistant")
            self.logger.error(f"Ошибка экспорта: {e}")

    def _cmd_import_folder(self, folder_path: str):
        if not folder_path:
            self.chat_widget.add_message("❌ Укажите путь к папке: /import_folder <путь>", "assistant")
            return
        folder_path = folder_path.strip().strip('"').strip("'")
        if not os.path.isdir(folder_path):
            self.chat_widget.add_message(f"❌ Папка не найдена: {folder_path}", "assistant")
            return
        supported_ext = {'.docx', '.xlsx', '.pdf', '.xls', '.txt', '.md', '.json', '.csv'}
        count = 0
        errors = []
        for ext in supported_ext:
            for file_path in glob.glob(os.path.join(folder_path, f"*{ext}")):
                try:
                    result = FileLoader._read_file(file_path, self.config_manager)
                    if result:
                        filename, content, file_format = result
                        self.dialog_manager.add_loaded_document(filename, content, file_format, file_path)
                        self.attached_files.append((filename, content))
                        count += 1
                except Exception as e:
                    errors.append(f"{Path(file_path).name}: {e}")
        self._update_file_label()
        if count > 0:
            self.chat_widget.add_message(f"✅ Загружено {count} документов из папки", "assistant")
            self.dialog_manager.add_message("assistant", f"Загружено {count} документов из папки")
        if errors:
            self.chat_widget.add_message(f"⚠️ Ошибки при загрузке:\n" + "\n".join(errors), "assistant")

    def _cmd_clear_documents(self):
        self.dialog_manager.clear_loaded_documents()
        self.attached_files = []
        self._update_file_label()
        self.chat_widget.add_message("✅ Загруженные документы очищены", "assistant")
        self.dialog_manager.add_message("assistant", "Загруженные документы очищены")

    def _cmd_clear_scenarios(self):
        self.dialog_manager.clear_loaded_scenarios()
        self.chat_widget.add_message("✅ Загруженные сценарии очищены", "assistant")
        self.dialog_manager.add_message("assistant", "Загруженные сценарии очищены")

    def _cmd_stop_scenario(self):
        """Останавливает выполняющийся сценарий (если есть)."""
        if self.analysis_thread and self.analysis_thread.is_alive():
            if self.cancel_event:
                self.cancel_event.set()
                self.chat_widget.add_message("⏹️ Запрос на остановку сценария отправлен...", "assistant")
                self.dialog_manager.add_message("assistant", "Запрос на остановку сценария отправлен...")
                self.statusbar.config(text="Остановка сценария...")
            else:
                self.chat_widget.add_message("⚠️ Сценарий выполняется, но событие отмены не найдено.", "assistant")
        else:
            self.chat_widget.add_message("ℹ️ Нет выполняющегося сценария для остановки.", "assistant")

    # ---------- Диалоги выбора файлов (меню) ----------
    def _cmd_load_document_dialog(self):
        file_path = filedialog.askopenfilename(
            title="Выберите документ для загрузки",
            filetypes=[
                ("Поддерживаемые", "*.docx *.xlsx *.pdf *.xls *.txt *.md *.json *.csv"),
                ("Документы Word", "*.docx"),
                ("Таблицы Excel", "*.xlsx *.xls"),
                ("PDF", "*.pdf"),
                ("Текстовые", "*.txt *.md *.json *.csv"),
                ("Все файлы", "*.*")
            ]
        )
        if file_path:
            self._load_single_document(file_path)

    def _cmd_load_scenario_dialog(self):
        file_path = filedialog.askopenfilename(
            title="Выберите JSON-сценарий",
            filetypes=[("JSON files", "*.json")]
        )
        if file_path:
            self._cmd_load_scenario(file_path)

    def _cmd_import_folder_dialog(self):
        folder_path = filedialog.askdirectory(title="Выберите папку с документами")
        if folder_path:
            self._cmd_import_folder(folder_path)

    # ---------- Загрузка/сохранение промптов (глобальные) ----------
    def _load_prompt_global(self):
        self.chat_widget._load_prompt_from_dialog()

    def _save_prompt_global(self):
        self.chat_widget._save_prompt_to_dialog()

    def _auto_load_prompt(self):
        prompt_defaults = self.config_manager.get_prompt_defaults()
        auto_file = prompt_defaults.get("auto_load_file", "")
        if auto_file and os.path.exists(auto_file):
            try:
                with open(auto_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.chat_widget.input_field.delete("1.0", END)
                self.chat_widget.input_field.insert("1.0", content)
                self.add_system_message(f"📂 Автозагрузка промпта из {auto_file}")
            except Exception as e:
                self.logger.error(f"Ошибка автозагрузки промпта: {e}")

    # ---------- Работа с файлами (прикрепление, очистка, метки) ----------
    def load_file(self):
        """Загружает файл через диалог и добавляет его в контекст чата"""
        result = FileLoader.load_from_dialog(self, self.config_manager)
        if result:
            filename, content, file_format = result
            self.attached_files.append((filename, content))
            self.dialog_manager.add_attached_file(filename, content)
            file_path = self._find_file_by_name(filename)
            self.dialog_manager.add_loaded_document(filename, content, file_format, file_path)
            self._update_file_label()
            self.statusbar.config(text=f"Файл '{filename}' добавлен")
            self.chat_widget.add_message(f"*Система*: Загружен файл \"{filename}\".", "assistant")
            saved_path = self._save_converted_file(filename, content, file_format)
            if saved_path:
                self.chat_widget.add_message(f"*Система*: Извлечённый текст сохранён в {saved_path}", "assistant")
        else:
            messagebox.showerror("Ошибка", "Не удалось загрузить файл.\nПоддерживаются: docx, xlsx, pdf, txt, md, py, json, csv")

    def _find_file_by_name(self, filename: str) -> str:
        if os.path.exists(filename):
            return filename
        docs_dir = self.config_manager.config.get("input_settings", {}).get("default_input_directory", "./documents")
        full_path = os.path.join(docs_dir, filename)
        if os.path.exists(full_path):
            return full_path
        converted_dir = self.config_manager.get_paths().get("converted_files_dir", "./converted_files")
        full_path = os.path.join(converted_dir, filename)
        if os.path.exists(full_path):
            return full_path
        return filename

    def _update_file_label(self):
        if self.attached_files:
            names = ", ".join(f[0] for f in self.attached_files)
            self.file_label.config(text=f"📎 Файлы: {names}")
            self.clear_file_btn.config(state=NORMAL)
        else:
            self.file_label.config(text="Файлы не загружены")
            self.clear_file_btn.config(state=DISABLED)

    def clear_file(self):
        self.attached_files = []
        self.dialog_manager.clear_attached_files()
        self._update_file_label()
        self.statusbar.config(text="Все файлы удалены")
        self.chat_widget.add_message("*Система*: Все загруженные файлы удалены.", "assistant")

    def _save_converted_file(self, original_filename: str, content: str, file_format: str) -> Optional[str]:
        try:
            save_dir = self.config_manager.get_paths().get("converted_files_dir", "./converted_files")
            Path(save_dir).mkdir(parents=True, exist_ok=True)
            base = Path(original_filename).stem
            ext = ".json" if file_format == "json" else ".md"
            save_path = Path(save_dir) / f"{base}{ext}"
            counter = 1
            while save_path.exists():
                save_path = Path(save_dir) / f"{base}_{counter}{ext}"
                counter += 1
            with open(save_path, 'w', encoding='utf-8') as f:
                f.write(content)
            self.logger.info(f"Сохранён конвертированный файл: {save_path}")
            return str(save_path)
        except Exception as e:
            self.logger.error(f"Ошибка сохранения конвертированного файла: {e}")
            return None

    # ---------- Работа с диалогами (сохранение/загрузка/очистка) ----------
    def new_dialog(self):
        if self.dialog_manager.current_dialog["messages"]:
            reply = messagebox.askyesno("Новый диалог", "Текущий диалог будет потерян. Продолжить?")
            if not reply:
                return
        self.dialog_manager.new_dialog()
        self.chat_widget.clear()
        self.attached_files = []
        self._update_file_label()
        self.statusbar.config(text="Новый диалог создан")

    def save_dialog(self):
        if not self.dialog_manager.current_dialog["messages"]:
            messagebox.showwarning("Сохранение", "Нет сообщений для сохранения.")
            return
        for fname, content in self.attached_files:
            self.dialog_manager.add_attached_file(fname, content)
        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json")],
            initialdir=self.dialog_manager.dialogs_dir,
            initialfile=f"{self.dialog_manager.current_dialog['dialog_id']}.json"
        )
        if file_path:
            self.dialog_manager.save_dialog(file_path)
            self.statusbar.config(text=f"Диалог сохранён в {file_path}")

    def load_dialog(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("JSON files", "*.json")],
            initialdir=self.dialog_manager.dialogs_dir
        )
        if file_path:
            self.dialog_manager.load_dialog(file_path)
            self.chat_widget.clear()
            for msg in self.dialog_manager.current_dialog["messages"]:
                self.chat_widget.add_message(msg["content"], msg["role"])
            self.attached_files = [(f["filename"], f["content"]) for f in
                                   self.dialog_manager.current_dialog.get("attached_files", [])]
            self._update_file_label()
            self.statusbar.config(text=f"Диалог загружен: {file_path}")

    def summarize_dialog(self):
        if not self.dialog_manager.current_dialog["messages"]:
            messagebox.showwarning("Обобщение", "Нет сообщений для обобщения.")
            return
        self.statusbar.config(text="Генерация summary...")
        dialog_params = self.config_manager.get_dialog_params()
        summary_pairs = dialog_params.get("summary_pairs", 3)
        pairs = self.dialog_manager.get_conversation_pairs(summary_pairs)
        conv_text = "Предыдущее обобщение:\n" + self.dialog_manager.current_dialog.get("summary", "") + "\n\n"
        for i, (q, a) in enumerate(pairs):
            conv_text += f"Вопрос {i+1}: {q}\nОтвет {i+1}: {a}\n"
        if self.attached_files:
            conv_text += "\nСодержимое загруженных файлов:\n"
            for fname, content in self.attached_files:
                conv_text += f"--- {fname} ---\n{content}\n"
        def generate():
            summary = self.chat_client.generate_summary(conv_text)
            self.after(0, lambda: self._on_summary_ready(summary))
        threading.Thread(target=generate, daemon=True).start()

    def _on_summary_ready(self, summary):
        self.dialog_manager.current_dialog["summary"] = summary
        self.dialog_manager.save_dialog()
        self.chat_widget.add_message(f"*Система*: Обобщение диалога:\n{summary}", "assistant")
        self.statusbar.config(text="Summary готов")

    # ---------- Интерфейс вкладки "Анализ документа" ----------
    def _select_document(self):
        file_paths = filedialog.askopenfilenames(
            title="Выберите один или несколько документов",
            filetypes=[
                ("Поддерживаемые", "*.docx *.xlsx *.pdf *.xls *.txt *.md"),
                ("Документы Word", "*.docx"),
                ("Таблицы Excel", "*.xlsx *.xls"),
                ("PDF", "*.pdf"),
                ("Текстовые", "*.txt *.md"),
                ("Все файлы", "*.*")
            ]
        )
        if file_paths:
            self.doc_paths_var = list(file_paths)
            self.doc_path_var.set("")
            self.doc_entry.config(state=NORMAL)
            self.doc_entry.delete(0, END)
            self.doc_entry.insert(0, f"Выбрано {len(file_paths)} файлов")
            self.doc_entry.config(state="readonly")
            self._update_analyze_button_state()

    def _select_scenario(self):
        file_path = filedialog.askopenfilename(
            title="Выберите JSON-сценарий",
            filetypes=[("JSON files", "*.json")]
        )
        if file_path:
            self.scenario_path_var.set(file_path)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    if 'stages' not in data:
                        raise ValueError("Нет ключа 'stages'")
                self._update_analyze_button_state()
            except Exception as e:
                messagebox.showerror("Ошибка", f"Невалидный сценарий:\n{e}")
                self.scenario_path_var.set("")

    def _update_analyze_button_state(self):
        enabled = bool(self.doc_path_var.get() and self.scenario_path_var.get())
        self.analyze_btn.config(state=NORMAL if enabled else DISABLED)

    def _open_scenario_editor(self):
        launch_scenario_editor(self, self.config_manager)

    def _start_analysis(self):
        if self.analysis_thread and self.analysis_thread.is_alive():
            messagebox.showwarning("Анализ", "Анализ уже выполняется")
            return
        # Определяем список файлов
        if hasattr(self, 'doc_paths_var') and self.doc_paths_var:
            doc_paths = self.doc_paths_var
        else:
            doc_path = self.doc_path_var.get()
            if doc_path:
                doc_paths = [doc_path]
            else:
                doc_paths = []
        if not doc_paths:
            messagebox.showwarning("Анализ", "Выберите документ(ы) для анализа")
            return
        scenario_path = self.scenario_path_var.get()
        if not scenario_path:
            messagebox.showwarning("Анализ", "Выберите сценарий")
            return
        self.analyze_btn.config(state=DISABLED)
        self.cancel_btn.config(state=NORMAL)
        self.progress_var.set(0)
        self.log_text.config(state=NORMAL)
        self.log_text.delete("1.0", END)
        self.log_text.insert(END, "Начало анализа...\n")
        self.log_text.config(state=DISABLED)
        self.result_text.config(state=NORMAL)
        self.result_text.delete("1.0", END)
        self.result_text.config(state=DISABLED)
        self._intermediate_results = {}
        self.intermediate_combo['values'] = []
        self.intermediate_combo.set("")
        self.intermediate_text.config(state=NORMAL)
        self.intermediate_text.delete("1.0", END)
        self.intermediate_text.config(state=DISABLED)
        api_key = self.config_manager.get_api_key()
        if not api_key:
            self._append_log("❌ API ключ не найден. Задайте его в настройках.")
            self._end_analysis(False)
            return
        llm_params = self.config_manager.get_llm_params()
        api_settings = {
            "api_key": api_key,
            "model": llm_params.get("model", "gpt-4o-mini"),
            "temperature": llm_params.get("temperature", 0.2),
            "max_tokens": llm_params.get("max_tokens", 2000),
            "timeout": llm_params.get("timeout", 60)
        }
        self.cancel_event = threading.Event()
        self.analysis_thread = threading.Thread(
            target=self._run_analysis_thread,
            args=(doc_paths, scenario_path, api_settings),
            daemon=True
        )
        self.analysis_thread.start()

    def _run_analysis_thread(self, doc_paths, scenario_path, api_settings):
        def progress_callback(msg_type, data):
            self.after(0, lambda: self._handle_progress(msg_type, data))

        try:
            result = run_analysis(
                document_paths=doc_paths,  # теперь список
                scenario_path=scenario_path,
                api_settings=api_settings,
                progress_callback=progress_callback,
                cancel_event=self.cancel_event
            )
            self.after(0, lambda: self._on_analysis_finished(result))
        except Exception as e:
            self.after(0, lambda: self._on_analysis_error(str(e)))

    def _handle_progress(self, msg_type, data):
        if msg_type == "log":
            self._append_log(data)
        elif msg_type == "stage":
            self._update_progress(data["current"], data["total"], data["name"])
        elif msg_type == "stage_result":
            stage_id = data.get("stage_id", "unknown")
            # Используем formatted_content, если он есть, иначе full_result
            formatted = data.get("formatted_content")
            if formatted:
                preview = formatted[:500] + ("..." if len(formatted) > 500 else "")
                self._append_log(f"--- Результат этапа '{stage_id}': ---")
                self._append_log(preview)
                self._append_log("---")
                self._intermediate_results[stage_id] = formatted
            else:
                preview = data.get("result_preview", "")
                self._append_log(f"--- Результат этапа '{stage_id}': ---")
                self._append_log(preview)
                self._append_log("---")
                self._intermediate_results[stage_id] = data.get("full_result", preview)
            self.after(0, self._update_intermediate_combo)

    def _update_progress(self, current, total, stage_name):
        percent = int((current / total) * 100)
        self.progress_var.set(percent)
        self._append_log(f"Этап {current}/{total}: {stage_name}")

    def _append_log(self, message):
        self.log_text.config(state=NORMAL)
        self.log_text.insert(END, f"{message}\n")
        self.log_text.see(END)
        self.log_text.config(state=DISABLED)

    def _on_analysis_finished(self, result):
        self._end_analysis(True)
        if result["status"] == "success":
            self._append_log("✅ Анализ завершён успешно")
            self._show_result(result["result"], result.get("format", "text"))
        elif result["status"] == "cancelled":
            self._append_log("⚠️ Анализ отменён пользователем")
        else:
            self._append_log(f"❌ Ошибка: {result.get('message', 'Неизвестная ошибка')}")

    def _on_analysis_error(self, error_msg):
        self._end_analysis(False)
        self._append_log(f"❌ Критическая ошибка: {error_msg}")

    def _end_analysis(self, success):
        self.after(0, lambda: self.analyze_btn.config(
            state=NORMAL if self.doc_path_var.get() and self.scenario_path_var.get() else DISABLED))
        self.after(0, lambda: self.cancel_btn.config(state=DISABLED))
        self.analysis_thread = None
        self.cancel_event = None

    def _cancel_analysis(self):
        if self.cancel_event:
            self.cancel_event.set()
            self._append_log("Отмена анализа...")

    def _show_result(self, text, fmt):
        self.result_text.config(state=NORMAL)
        self.result_text.delete("1.0", END)
        self.result_text.insert(END, text)
        self.result_text.config(state=DISABLED)
        self.result_format = fmt

    def _copy_result(self):
        text = self.result_text.get("1.0", END).strip()
        if text:
            self.clipboard_clear()
            self.clipboard_append(text)
            self.statusbar.config(text="Результат скопирован в буфер обмена")

    def _save_result(self):
        text = self.result_text.get("1.0", END).strip()
        if not text:
            return
        fmt = getattr(self, 'result_format', 'text')

        filetypes = [
            ("Текстовый файл", "*.txt"),
            ("Markdown", "*.md"),
            ("JSON", "*.json"),
            ("CSV", "*.csv"),
            ("Документ Word", "*.docx"),
            ("Таблица Excel", "*.xlsx")
        ]
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=filetypes)
        if not file_path:
            return

        ext = Path(file_path).suffix.lower()
        try:
            if ext == '.docx':
                self._save_as_docx(file_path, text, fmt)
            elif ext == '.xlsx':
                self._save_as_xlsx(file_path, text, fmt)
            elif ext == '.json':
                # Сохраняем JSON с корректной кодировкой
                try:
                    data = json.loads(text)
                    with open(file_path, 'w', encoding='utf-8') as f:
                        json.dump(data, f, ensure_ascii=False, indent=2)
                except json.JSONDecodeError:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                self.statusbar.config(text=f"Результат сохранён: {file_path}")
            else:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                self.statusbar.config(text=f"Результат сохранён: {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка сохранения", f"Не удалось сохранить файл:\n{e}")
            self.logger.error(f"Ошибка сохранения: {e}")

    def _on_intermediate_selected(self, event):
        stage_id = self.intermediate_combo.get()
        if hasattr(self, '_intermediate_results') and stage_id in self._intermediate_results:
            self.intermediate_text.config(state=NORMAL)
            self.intermediate_text.delete("1.0", END)
            self.intermediate_text.insert(END, self._intermediate_results[stage_id])
            self.intermediate_text.config(state=DISABLED)

    def _update_intermediate_combo(self):
        if hasattr(self, '_intermediate_results'):
            stages = list(self._intermediate_results.keys())
            current = self.intermediate_combo.get()
            self.intermediate_combo['values'] = stages
            if current in stages:
                self.intermediate_combo.set(current)
            elif stages:
                self.intermediate_combo.set(stages[-1])
                self._on_intermediate_selected(None)

    # ---------- Сохранение в DOCX и XLSX (для результатов и чата) ----------
    def _parse_markdown_table(self, table_lines):
        lines = [line.strip() for line in table_lines if line.strip()]
        if len(lines) < 2:
            return None
        header_line = lines[0]
        sep_line = lines[1] if len(lines) > 1 else None
        if not sep_line or '---' not in sep_line:
            data_rows = lines
            headers = []
        else:
            headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
            data_rows = lines[2:]
        table = []
        if headers:
            table.append(headers)
        for row_line in data_rows:
            cells = [cell.strip() for cell in row_line.split('|') if cell.strip()]
            if len(cells) < len(headers):
                cells.extend([''] * (len(headers) - len(cells)))
            table.append(cells[:len(headers)])
        return table

    def _convert_csv_to_table(self, text: str) -> list:
        """
        Преобразует CSV-текст (строки с разделителями) в список списков для таблицы.
        Определяет разделитель автоматически (запятая, табуляция, точка с запятой).
        """
        if not text:
            return None

        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if len(lines) < 2:
            return None

        # Определяем разделитель
        separators = [',', '\t', ';', '|']
        best_sep = None
        best_count = 0

        for sep in separators:
            count = sum(line.count(sep) for line in lines)
            if count > best_count:
                best_count = count
                best_sep = sep

        if best_sep is None or best_count < 2:
            return None

        table = []
        for line in lines:
            # Разбиваем строку, учитывая кавычки
            if best_sep == ',':
                import csv
                import io
                reader = csv.reader(io.StringIO(line))
                row = next(reader)
            else:
                row = [cell.strip() for cell in line.split(best_sep)]
            table.append(row)

        return table

    def _is_likely_csv(self, text: str) -> bool:
        """Определяет, похож ли текст на CSV."""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if len(lines) < 2:
            return False

        # Проверяем наличие разделителей в большинстве строк
        separators = [',', '\t', ';', '|']
        for sep in separators:
            count = sum(1 for line in lines if sep in line)
            if count >= len(lines) * 0.5:  # хотя бы 50% строк содержат разделитель
                return True
        return False

    def _convert_csv_block_to_table(self, csv_text: str) -> list:
        """
        Преобразует CSV-текст (с кавычками) в список списков для таблицы.
        """
        if not csv_text:
            return None

        import csv
        import io

        lines = csv_text.strip().split('\n')
        if len(lines) < 2:
            return None

        try:
            # Пробуем прочитать как CSV с кавычками
            reader = csv.reader(io.StringIO(csv_text))
            table = list(reader)
            if table and len(table) > 0 and len(table[0]) > 0:
                return table
        except:
            pass

        # Если не удалось, пробуем другие разделители
        separators = ['\t', ';', '|']
        for sep in separators:
            try:
                table = [line.split(sep) for line in lines]
                if table and len(table) > 0 and len(table[0]) > 1:
                    return table
            except:
                continue

        return None

    def _save_as_docx(self, file_path: str, content: str, fmt: str):
        """Сохраняет результат в формате DOCX с поддержкой заголовков, жирного/курсива и CSV-блоков."""
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import re
        import csv
        import io

        doc = Document()

        if fmt == 'markdown':
            lines = content.split('\n')
            in_table = False
            table_lines = []
            in_code_block = False
            code_block_lines = []
            code_block_lang = ''
            csv_buffer = []
            in_csv = False

            def flush_csv():
                nonlocal csv_buffer, in_csv
                if csv_buffer:
                    table = self._convert_csv_block_to_table('\n'.join(csv_buffer))
                    if table and len(table) > 0 and len(table[0]) > 0:
                        word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                        word_table.style = 'Table Grid'
                        for r, row_data in enumerate(table):
                            for c, cell_text in enumerate(row_data):
                                if c < len(row_data):
                                    word_table.cell(r, c).text = cell_text
                    else:
                        p = doc.add_paragraph()
                        run = p.add_run('\n'.join(csv_buffer))
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    csv_buffer = []
                    in_csv = False

            def add_paragraph_with_markdown(text: str):
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
                for part in parts:
                    if not part:
                        continue
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                    else:
                        p.add_run(part)

            for line in lines:
                # Пропускаем пустые строки, но не в CSV и не в коде
                if not line:
                    if in_code_block:
                        code_block_lines.append('')
                    elif in_csv:
                        csv_buffer.append('')
                    else:
                        doc.add_paragraph()
                    continue

                # ---- Обработка кодовых блоков ----
                if line.startswith('```'):
                    flush_csv()
                    if in_code_block:
                        # Закрываем блок
                        if code_block_lines:
                            block_text = '\n'.join(code_block_lines)
                            if code_block_lang == 'csv':
                                table = self._convert_csv_block_to_table(block_text)
                                if table and len(table) > 0 and len(table[0]) > 0:
                                    word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                                    word_table.style = 'Table Grid'
                                    for r, row_data in enumerate(table):
                                        for c, cell_text in enumerate(row_data):
                                            if c < len(row_data):
                                                word_table.cell(r, c).text = cell_text
                                else:
                                    p = doc.add_paragraph()
                                    run = p.add_run(block_text)
                                    run.font.name = 'Courier New'
                                    run.font.size = Pt(9)
                            else:
                                p = doc.add_paragraph()
                                run = p.add_run(block_text)
                                run.font.name = 'Courier New'
                                run.font.size = Pt(9)
                            code_block_lines = []
                            code_block_lang = ''
                        in_code_block = False
                    else:
                        in_code_block = True
                        code_block_lang = line[3:].strip().lower()
                    continue

                if in_code_block:
                    code_block_lines.append(line)
                    continue

                # ---- Проверка на CSV (если строка похожа на CSV и не является заголовком) ----
                if not line.startswith('#') and '|' not in line and self._is_likely_csv(line):
                    if not in_csv:
                        flush_csv()
                        in_csv = True
                    csv_buffer.append(line)
                    continue
                elif in_csv:
                    flush_csv()

                # ---- Обработка Markdown-таблиц ----
                if '|' in line and not line.strip().startswith('|'):
                    if not in_table:
                        in_table = True
                    table_lines.append(line)
                else:
                    if in_table:
                        if table_lines:
                            table = self._parse_markdown_table(table_lines)
                            if table and len(table) > 0 and len(table[0]) > 0:
                                word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                                word_table.style = 'Table Grid'
                                for r, row_data in enumerate(table):
                                    for c, cell_text in enumerate(row_data):
                                        if c < len(row_data):
                                            word_table.cell(r, c).text = cell_text
                        table_lines = []
                        in_table = False

                    # ---- Обработка заголовков Markdown ----
                    header_match = re.match(r'^(#+)\s+(.*)$', line)
                    if header_match:
                        level = len(header_match.group(1))
                        text = header_match.group(2).strip()
                        # Удаляем ** если они есть (чтобы не дублировать жирность)
                        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
                        p = doc.add_paragraph()
                        run = p.add_run(text)
                        run.bold = True
                        if level == 1:
                            run.font.size = Pt(18)
                            p.style = 'Heading 1'
                        elif level == 2:
                            run.font.size = Pt(16)
                            p.style = 'Heading 2'
                        elif level == 3:
                            run.font.size = Pt(14)
                            p.style = 'Heading 3'
                        elif level >= 4:
                            run.font.size = Pt(12)
                            p.style = 'Heading 4'
                        continue

                    # ---- Обычный текст с Markdown ----
                    add_paragraph_with_markdown(line)

            # ---- Обработка остатков ----
            flush_csv()
            if in_table and table_lines:
                table = self._parse_markdown_table(table_lines)
                if table and len(table) > 0 and len(table[0]) > 0:
                    word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                    word_table.style = 'Table Grid'
                    for r, row_data in enumerate(table):
                        for c, cell_text in enumerate(row_data):
                            if c < len(row_data):
                                word_table.cell(r, c).text = cell_text

        elif fmt == 'json':
            try:
                # Проверяем, является ли содержимое валидным JSON
                if isinstance(content, str):
                    data = json.loads(content)
                else:
                    data = content

                # Если JSON — массив объектов, создаём таблицу
                if isinstance(data, list) and data and all(isinstance(item, dict) for item in data):
                    self._json_to_word_table(doc, data)
                # Если JSON — объект, создаём таблицу ключ-значение
                elif isinstance(data, dict):
                    self._json_to_word_table(doc, data)
                else:
                    # Иначе — вставляем как код с правильной кодировкой
                    p = doc.add_paragraph()
                    run = p.add_run(json.dumps(data, ensure_ascii=False, indent=2))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            except (json.JSONDecodeError, TypeError) as e:
                self.logger.warning(f"Не удалось распарсить JSON: {e}")
                # Вставляем как текст с правильной кодировкой
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        else:
            # CSV, text – вставляем моноширинным шрифтом
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

        doc.save(file_path)
        self.statusbar.config(text=f"Результат сохранён в DOCX: {file_path}")

    def _save_as_xlsx(self, file_path: str, content: str, fmt: str):
        wb = Workbook()
        if fmt in ('json', 'csv'):
            try:
                if fmt == 'json':
                    data = json.loads(content)
                    # Если данные — это словарь с ключом 'systems', извлекаем массив
                    if isinstance(data, dict) and 'systems' in data:
                        data = data['systems']
                    if isinstance(data, list):
                        df = pd.DataFrame(data)
                    else:
                        df = pd.DataFrame([data])
                else:  # csv
                    df = pd.read_csv(io.StringIO(content))
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False, sheet_name='Data')
                self.statusbar.config(text=f"Результат сохранён в XLSX: {file_path}")
                return
            except Exception as e:
                self.logger.warning(f"Не удалось преобразовать в таблицу: {e}")
                # Продолжаем как обычный текст

        if fmt == 'markdown':
            lines = content.split('\n')
            table_lines = []
            in_table = False
            text_parts = []
            tables = []
            for line in lines:
                if '|' in line and not line.strip().startswith('|'):
                    if not in_table:
                        in_table = True
                    table_lines.append(line)
                else:
                    if in_table:
                        if table_lines:
                            table = self._parse_markdown_table(table_lines)
                            if table:
                                tables.append(table)
                        table_lines = []
                        in_table = False
                        text_parts.append(line)
                    else:
                        text_parts.append(line)
            if table_lines:
                table = self._parse_markdown_table(table_lines)
                if table:
                    tables.append(table)

            if text_parts:
                ws_text = wb.active
                ws_text.title = "Text"
                ws_text.cell(row=1, column=1, value='\n'.join(text_parts))

            if tables:
                ws_table = wb.create_sheet("Tables")
                start_row = 1
                for table in tables:
                    for row_data in table:
                        for col_idx, cell_value in enumerate(row_data, 1):
                            ws_table.cell(row=start_row, column=col_idx, value=cell_value)
                        start_row += 1
                    start_row += 1
        else:
            ws = wb.active
            ws.cell(row=1, column=1, value=content)

        wb.save(file_path)
        self.statusbar.config(text=f"Результат сохранён в XLSX: {file_path}")

    def _json_to_word_table(self, doc, data):
        """Преобразует JSON-данные в таблицу Word с корректной обработкой длинных строк."""
        from docx.shared import Inches

        if isinstance(data, list) and data and all(isinstance(item, dict) for item in data):
            headers = list(data[0].keys())
            table = doc.add_table(rows=1 + len(data), cols=len(headers))
            table.style = 'Table Grid'
            # Заголовки
            for col_idx, header in enumerate(headers):
                cell = table.cell(0, col_idx)
                cell.text = header
                # Жирный шрифт для заголовков
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            # Данные
            for row_idx, item in enumerate(data, 1):
                for col_idx, key in enumerate(headers):
                    val = item.get(key, '')
                    if isinstance(val, (list, dict)):
                        val = json.dumps(val, ensure_ascii=False)
                    cell = table.cell(row_idx, col_idx)
                    cell.text = str(val)
                    # Автоподбор ширины столбцов
                    cell.width = None
        elif isinstance(data, dict):
            table = doc.add_table(rows=len(data), cols=2)
            table.style = 'Table Grid'
            for row_idx, (key, val) in enumerate(data.items()):
                table.cell(row_idx, 0).text = str(key)
                if isinstance(val, (list, dict)):
                    val = json.dumps(val, ensure_ascii=False)
                table.cell(row_idx, 1).text = str(val)
        else:
            p = doc.add_paragraph()
            run = p.add_run(json.dumps(data, ensure_ascii=False))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

    def _save_chat_as_docx_wrapper(self):
        chat_text = self.chat_widget.get_chat_text()
        if not chat_text:
            messagebox.showwarning("Сохранение", "Чат пуст, нечего сохранять")
            return
        self._save_chat_as_docx(chat_text)

    def _save_chat_as_docx(self, chat_text: str):
        """Сохраняет содержимое чата в DOCX с поддержкой Markdown-разметки, заголовков и CSV-таблиц."""
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        doc = Document()

        # Парсим текст чата построчно
        lines = chat_text.split('\n')
        in_code_block = False
        code_block_lines = []
        code_block_lang = ''
        csv_buffer = []
        in_csv = False

        def flush_csv():
            nonlocal csv_buffer, in_csv
            if csv_buffer:
                table = self._convert_csv_to_table('\n'.join(csv_buffer))
                if table and len(table) > 0 and len(table[0]) > 0:
                    word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                    word_table.style = 'Table Grid'
                    for i, row_data in enumerate(table):
                        for j, cell_text in enumerate(row_data):
                            if j < len(row_data):
                                word_table.cell(i, j).text = cell_text
                else:
                    # Если не удалось преобразовать в таблицу, вставляем как текст
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(csv_buffer))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                csv_buffer = []
                in_csv = False

        def add_paragraph_with_markdown(text: str):
            """Добавляет параграф с обработкой Markdown-разметки (жирный, курсив, код)."""
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                else:
                    p.add_run(part)

        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            i += 1

            # Пропускаем пустые строки (но не в CSV-блоке)
            if not line:
                if in_csv:
                    csv_buffer.append('')
                else:
                    doc.add_paragraph()
                continue

            # Пропускаем разделители (---)
            if line.startswith('---') and line.endswith('---'):
                continue

            # Обработка маркеров чата (Вы, Ассистент, Система)
            if line.startswith('Вы  ') or line.startswith('Ассистент  ') or line.startswith('*Система*'):
                flush_csv()
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(10)
                continue

            # Обработка кодовых блоков
            if line.startswith('```'):
                flush_csv()
                if in_code_block:
                    # Закрываем блок
                    if code_block_lines:
                        # Проверяем, является ли блок CSV
                        block_text = '\n'.join(code_block_lines)
                        if code_block_lang == 'csv':
                            table = self._convert_csv_block_to_table(block_text)
                            if table and len(table) > 0 and len(table[0]) > 0:
                                word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                                word_table.style = 'Table Grid'
                                for r, row_data in enumerate(table):
                                    for c, cell_text in enumerate(row_data):
                                        if c < len(row_data):
                                            word_table.cell(r, c).text = cell_text
                            else:
                                # Если не удалось преобразовать в таблицу, вставляем как код
                                p = doc.add_paragraph()
                                run = p.add_run(block_text)
                                run.font.name = 'Courier New'
                                run.font.size = Pt(9)
                        else:
                            p = doc.add_paragraph()
                            run = p.add_run(block_text)
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                        code_block_lines = []
                        code_block_lang = ''
                    in_code_block = False
                else:
                    in_code_block = True
                    # Определяем язык блока
                    lang = line[3:].strip().lower()
                    code_block_lang = lang
                continue

            if in_code_block:
                code_block_lines.append(line)
                continue

            # Проверка на CSV (если строка похожа на CSV и не является заголовком)
            if not line.startswith('#') and self._is_likely_csv(line):
                if not in_csv:
                    flush_csv()
                    in_csv = True
                csv_buffer.append(line)
                continue
            elif in_csv:
                flush_csv()

            # Обработка Markdown-заголовков
            header_match = re.match(r'^(#+)\s+(.*)$', line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2).strip()
                # Удаляем ** если они есть
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)

                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                if level == 1:
                    run.font.size = Pt(18)
                    p.style = 'Heading 1'
                elif level == 2:
                    run.font.size = Pt(16)
                    p.style = 'Heading 2'
                elif level == 3:
                    run.font.size = Pt(14)
                    p.style = 'Heading 3'
                elif level >= 4:
                    run.font.size = Pt(12)
                    p.style = 'Heading 4'
                continue

            # Обычный текст с Markdown
            flush_csv()
            add_paragraph_with_markdown(line)

        # Обработка остатков
        flush_csv()
        if in_code_block and code_block_lines:
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_block_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

        # Сохраняем файл
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Документ Word", "*.docx"), ("Все файлы", "*.*")],
            title="Сохранить чат как DOCX"
        )
        if file_path:
            doc.save(file_path)
            self.statusbar.config(text=f"Чат сохранён в DOCX: {file_path}")

    # ---------- Копирование чата ----------
    def copy_chat(self):
        chat_text = self.chat_widget.get_chat_text()
        if chat_text:
            self.clipboard_clear()
            self.clipboard_append(chat_text)
            self.statusbar.config(text="Чат скопирован в буфер обмена")
            self.chat_widget.add_message("*Система*: Содержимое чата скопировано в буфер обмена.", "assistant")
        else:
            self.statusbar.config(text="Чат пуст, нечего копировать")

    def clear_chat(self):
        self.chat_widget.clear()
        self.statusbar.config(text="Чат очищен")
        self.chat_widget.add_message("*Система*: История чата очищена.", "assistant")

    # ---------- Сохранение выделения из чата ----------
    def _safe_filename(self, title: str, default: str = "untitled") -> str:
        """Преобразует строку в безопасное имя файла (убирает недопустимые символы)."""
        import re
        # Убираем символы, запрещённые в именах файлов Windows/Linux
        safe = re.sub(r'[<>:"/\\|?*]', '_', title.strip())
        # Убираем лишние пробелы и точки в конце
        safe = safe.strip(' .')
        if not safe:
            safe = default
        # Ограничиваем длину
        if len(safe) > 100:
            safe = safe[:100]
        return safe

    def _save_selected_as_file(self):
        """Сохраняет выделенный текст чата в один файл с именем из первой строки."""
        selected = self.chat_widget.get_selected_text()
        if not selected:
            messagebox.showwarning("Сохранение", "Ничего не выделено в чате.")
            return

        # Первая строка (до первого переноса) как имя файла
        first_line = selected.splitlines()[0].strip() if selected else ""
        if not first_line:
            first_line = "selected_text"

        base_name = self._safe_filename(first_line, "selected_text")
        default_ext = ".md" if first_line.startswith('#') else ".txt"

        file_path = filedialog.asksaveasfilename(
            defaultextension=default_ext,
            filetypes=[("Текстовые файлы", "*.txt"), ("Markdown", "*.md"), ("Все файлы", "*.*")],
            initialfile=base_name + default_ext,
            title="Сохранить выделенный текст"
        )
        if not file_path:
            return

        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(selected)
            self.statusbar.config(text=f"Выделенный текст сохранён в {file_path}")
            self.chat_widget.add_message(f"*Система*: Выделенный текст сохранён в {file_path}", "assistant")
        except Exception as e:
            messagebox.showerror("Ошибка сохранения", f"Не удалось сохранить файл:\n{e}")

        def _save_selected_as_sections(self):
            """
            Сохраняет выделенный текст, разделяя по строкам, начинающимся с "---".
            Каждый раздел сохраняется в отдельный файл.
            Имя файла берётся из строки "--- <имя_файла>".
            Удаляет обратные кавычки (```) и возможный язык (например, json) в начале и конце содержимого.
            """
            selected = self.chat_widget.get_selected_text()
            if not selected:
                messagebox.showwarning("Сохранение", "Ничего не выделено в чате.")
                return

            import re

            # Разбиваем по строкам, начинающимся с "---" (возможны пробелы перед)
            pattern = r'(?=^[ \t]*---\s+.*$)'
            sections = re.split(pattern, selected, flags=re.MULTILINE)
            sections = [s.strip() for s in sections if s.strip()]

            if not sections:
                messagebox.showwarning("Сохранение",
                                       "Выделенный текст не содержит маркеров разделов (строк, начинающихся с '---').")
                return

            if len(sections) == 1:
                reply = messagebox.askyesno("Сохранение",
                                            "Текст не разбит на разделы (только один маркер).\n"
                                            "Сохранить как один файл?")
                if reply:
                    self._save_selected_as_file()
                return

            folder_path = filedialog.askdirectory(title="Выберите папку для сохранения разделов")
            if not folder_path:
                return

            saved_count = 0
            errors = []

            for section in sections:
                lines = section.splitlines()
                if not lines:
                    continue

                # Первая строка должна начинаться с "---"
                first_line = lines[0].strip()
                if not first_line.startswith('---'):
                    title = "unnamed_section"
                    content = section
                else:
                    # Извлекаем имя файла после "---"
                    file_name_part = first_line[3:].strip()
                    file_name_part = file_name_part.strip('"').strip("'")
                    if file_name_part:
                        title = file_name_part
                    else:
                        title = "unnamed_section"
                    # Остальные строки (кроме первой) — содержимое
                    content = "\n".join(lines[1:]).strip()

                # Удаляем обратные кавычки и возможный язык в начале/конце
                # Если содержимое начинается с ``` и заканчивается ```
                if content.startswith('```') and content.endswith('```'):
                    lines = content.splitlines()
                    # Проверяем, что первая строка начинается с ``` (возможно с языком)
                    if lines and lines[0].strip().startswith('```'):
                        # Удаляем первую строку
                        lines = lines[1:]
                    if lines and lines[-1].strip() == '```':
                        # Удаляем последнюю строку
                        lines = lines[:-1]
                    content = "\n".join(lines).strip()

                # Дополнительно: если после удаления всё ещё есть обратные кавычки, удаляем их (на всякий случай)
                content = re.sub(r'^```[a-zA-Z]*\s*', '', content)
                content = re.sub(r'\s*```$', '', content)

                # Формируем имя файла
                base_name = self._safe_filename(title, "section")
                # Если имя не содержит расширения, добавляем .txt
                if '.' not in os.path.splitext(base_name)[1]:
                    base_name += ".txt"

                file_path = os.path.join(folder_path, base_name)

                counter = 1
                while os.path.exists(file_path):
                    name, ext = os.path.splitext(base_name)
                    file_path = os.path.join(folder_path, f"{name}_{counter}{ext}")
                    counter += 1

                try:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                    saved_count += 1
                except Exception as e:
                    errors.append(f"{file_path}: {e}")

            if saved_count > 0:
                self.statusbar.config(text=f"Сохранено {saved_count} разделов в {folder_path}")
                self.add_system_message(f"📁 Сохранено {saved_count} разделов в {folder_path}")
            if errors:
                self.add_system_message(f"⚠️ Ошибки при сохранении:\n" + "\n".join(errors))

    # ---------- Остальные обработчики (API, сигналы) ----------
    def _on_ai_chunk(self, chunk):
        self.after(0, lambda: self._append_chunk(chunk))

    def _append_chunk(self, chunk):
        self._stream_buffer += chunk

    def _on_ai_response(self, text, metadata):
        self.after(0, lambda: self._display_response(text))

    def _display_response(self, text):
        full = text if text else self._stream_buffer
        if full:
            self.chat_widget.add_message(full, "assistant")
            self.dialog_manager.add_message("assistant", full, tokens_used=0)
        else:
            self.chat_widget.add_message("*Ассистент не дал ответа*", "assistant")
        self._stream_buffer = ""

    def _on_api_error(self, msg, recoverable):
        self.after(0, lambda: self._show_error(msg, recoverable))

    def _show_error(self, msg, recoverable):
        self.logger.error(f"API error: {msg}")
        self.chat_widget.add_message(f"*Ошибка*: {msg}", "assistant")
        if not recoverable:
            self.statusbar.config(text="Критическая ошибка API. Проверьте ключ.")
        else:
            self.statusbar.config(text="Ошибка API, повторите попытку.")

    def _on_request_start(self):
        self.after(0, lambda: self._set_sending_state(True))

    def _on_request_finish(self):
        self.after(0, lambda: self._set_sending_state(False))

    def _set_sending_state(self, is_sending):
        if is_sending:
            self.chat_widget.set_send_enabled(False)
            self.statusbar.config(text="Ассистент печатает...")
        else:
            self.chat_widget.set_send_enabled(True)
            self.statusbar.config(text="Готов")

    # ---------- Настройки, проверка API, загрузка моделей ----------
    def open_settings(self):
        dialog = SettingsDialog(self, self.config_manager, self.available_models)
        self.wait_window(dialog)
        theme = self.config_manager.config.get("ui", {}).get("theme", "light")
        new_theme = "darkly" if theme == "dark" else "cosmo"
        self.style.theme_use(new_theme)

    def check_api_key(self):
        if not self.config_manager.get_api_key():
            reply = messagebox.askyesno("API ключ", "OPENAI_API_KEY не найден. Ввести сейчас?")
            if reply:
                from tkinter import simpledialog
                key = simpledialog.askstring("API ключ", "Введите OpenAI API ключ:", show='*')
                if key:
                    self.config_manager.set_api_key(key)
                    self.statusbar.config(text="API ключ сохранён")
                    self.chat_widget.add_message("*Система*: API ключ успешно сохранён.", "assistant")
                else:
                    self.statusbar.config(text="API ключ не задан")
            else:
                self.statusbar.config(text="API ключ не задан")

    def test_api_connection(self):
        def test():
            success = self.chat_client.test_connection()
            self.after(0, lambda: self._on_test_result(success))
        threading.Thread(target=test, daemon=True).start()

    def _on_test_result(self, success):
        if success:
            self.chat_widget.add_message("*Система*: ✅ Соединение с OpenAI API установлено.", "assistant")
            self.statusbar.config(text="API готов")
        else:
            self.chat_widget.add_message("*Система*: ❌ Не удалось подключиться к API. Проверьте ключ и интернет.", "assistant")
            self.statusbar.config(text="Ошибка подключения к API")

    def load_models_at_startup(self):
        def load():
            try:
                api_key = self.config_manager.get_api_key()
                if not api_key:
                    self.logger.warning("Не удалось загрузить модели: нет API ключа")
                    return
                from openai import OpenAI
                client = OpenAI(api_key=api_key)
                models = client.models.list()
                model_names = [m.id for m in models if 'gpt' in m.id and not m.id.startswith('whisper')]
                priority = ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo", "gpt-4"]
                sorted_models = sorted(model_names, key=lambda x: (x not in priority, priority.index(x) if x in priority else len(priority)))
                self.available_models = sorted_models
                self.logger.info(f"Загружено {len(sorted_models)} моделей")
            except Exception as e:
                self.logger.error(f"Ошибка загрузки моделей: {e}")
                self.available_models = ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo", "gpt-4"]
        threading.Thread(target=load, daemon=True).start()

    def show_about(self):
        messagebox.showinfo(
            "О программе",
            "AI Document Analyst (Chat)\nВерсия 1.0\n\n"
            "Агентная система анализа документов с ИИ.\n"
            "Использует OpenAI API для ответов на вопросы.\n\n© 2024"
        )

    # ---------- Рекомендация по разбиению файла ----------
    def _recommend_file_split(self, content: str, model: str = "gpt-4o-mini") -> str:
        try:
            import tiktoken
            encoding = tiktoken.encoding_for_model(model)
            total_tokens = len(encoding.encode(content))
            max_safe_tokens = 100000 if ("gpt-4" in model or "gpt-5" in model) else 50000
            if total_tokens <= max_safe_tokens:
                return None
            parts = (total_tokens // max_safe_tokens) + 1
            tokens_per_part = total_tokens // parts + 1
            return (
                f"⚠️ **Файл слишком большой для модели {model}.**\n"
                f"Размер: ~{total_tokens:,} токенов.\n"
                f"Рекомендация: разбейте файл на **{parts}** частей (примерно по {tokens_per_part:,} токенов каждая)\n"
                f"и загружайте их по очереди, задавая вопросы к каждой части отдельно.\n"
                f"Или используйте команду /load_document для загрузки отдельных файлов."
            )
        except Exception as e:
            self.logger.error(f"Ошибка подсчёта токенов: {e}")
            return None