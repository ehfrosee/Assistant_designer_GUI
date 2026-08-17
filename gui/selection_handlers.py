# -*- coding: utf-8 -*-
"""Сохранение выделенного текста из чата"""

import os
import re
from tkinter import filedialog, messagebox

from docx import Document
from docx.shared import Pt, Inches

from gui.export_handlers import ExportHandlers


class SelectionHandlers(ExportHandlers):
    """Обработчики сохранения выделенного текста."""

    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent

    def _safe_filename(self, title: str, default: str = "untitled") -> str:
        """Преобразует строку в безопасное имя файла."""
        safe = re.sub(r'[<>:"/\\|?*]', '_', title.strip())
        safe = safe.strip(' .')
        if not safe:
            safe = default
        if len(safe) > 100:
            safe = safe[:100]
        return safe

    def _detect_list_type(self, line: str) -> tuple:
        """Определяет тип списка и уровень вложенности."""
        stripped = line.lstrip()
        indent = len(line) - len(stripped)
        level = indent // 4
        ordered_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if ordered_match:
            return ('ordered', level, ordered_match.group(2), int(ordered_match.group(1)))
        if stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('• '):
            return ('unordered', level, stripped[2:].strip(), None)
        sub_ordered_match = re.match(r'^(\d+\.\d+)\s+(.*)$', stripped)
        if sub_ordered_match:
            return ('ordered', level + 1, sub_ordered_match.group(2), None)
        return (None, 0, stripped, None)

    def save_selected_as_file(self):
        """Сохраняет выделенный текст в один файл."""
        selected = self.parent.chat_widget.get_selected_text()
        if not selected:
            messagebox.showwarning("Сохранение", "Ничего не выделено в чате.")
            return
        first_line = selected.splitlines()[0].strip() if selected else ""
        if not first_line:
            first_line = "selected_text"
        base_name = self._safe_filename(first_line, "selected_text")
        default_ext = ".md" if first_line.startswith('#') else ".txt"
        file_path = filedialog.asksaveasfilename(
            defaultextension=default_ext,
            filetypes=[("Текстовые файлы", "*.txt"), ("Markdown", "*.md"), ("Все файлы", "*.*")],
            initialfile=base_name + default_ext,
            title="Сохранить выделенный текст"
        )
        if not file_path:
            return
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(selected)
            self.parent.statusbar.config(text=f"Выделенный текст сохранён в {file_path}")
            self.parent.add_system_message(f"📄 Выделенный текст сохранён в {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка сохранения", f"Не удалось сохранить файл:\n{e}")

    def save_selected_as_sections(self):
        """Сохраняет выделенный текст по разделам (маркер ---)."""
        selected = self.parent.chat_widget.get_selected_text()
        if not selected:
            messagebox.showwarning("Сохранение", "Ничего не выделено в чате.")
            return
        pattern = r'(?=^[ \t]*---\s+\S+.*$)'
        sections = re.split(pattern, selected, flags=re.MULTILINE)
        sections = [s.strip() for s in sections if s.strip()]
        if not sections:
            messagebox.showwarning("Сохранение", "Выделенный текст не содержит маркеров разделов (---).")
            return
        if len(sections) == 1:
            reply = messagebox.askyesno("Сохранение", "Текст не разбит на разделы. Сохранить как один файл?")
            if reply:
                self.save_selected_as_file()
            return
        folder_path = filedialog.askdirectory(title="Выберите папку для сохранения разделов")
        if not folder_path:
            return
        saved_count = 0
        errors = []
        for section in sections:
            lines = section.splitlines()
            if not lines:
                continue
            first_line = lines[0].strip()
            if not first_line.startswith('---'):
                continue
            file_name_part = first_line[3:].strip().strip('"').strip("'")
            if not file_name_part:
                continue
            title = file_name_part
            content = "\n".join(lines[1:]).strip()
            if content.startswith('```') and content.endswith('```'):
                lines_content = content.splitlines()
                if lines_content and lines_content[0].strip().startswith('```'):
                    lines_content = lines_content[1:]
                if lines_content and lines_content[-1].strip() == '```':
                    lines_content = lines_content[:-1]
                content = "\n".join(lines_content).strip()
            content = re.sub(r'^```[a-zA-Z]*\s*', '', content)
            content = re.sub(r'\s*```$', '', content)
            base_name = self._safe_filename(title, "section")
            if '.' not in os.path.splitext(base_name)[1]:
                base_name += ".txt"
            file_path = os.path.join(folder_path, base_name)
            counter = 1
            while os.path.exists(file_path):
                name, ext = os.path.splitext(base_name)
                file_path = os.path.join(folder_path, f"{name}_{counter}{ext}")
                counter += 1
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                saved_count += 1
            except Exception as e:
                errors.append(f"{file_path}: {e}")
        if saved_count > 0:
            self.parent.statusbar.config(text=f"Сохранено {saved_count} разделов в {folder_path}")
            self.parent.add_system_message(f"📁 Сохранено {saved_count} разделов в {folder_path}")
        if errors:
            self.parent.add_system_message(f"⚠️ Ошибки при сохранении:\n" + "\n".join(errors))

    def save_selected_as_docx(self):
        """Сохраняет выделенный текст в DOCX с поддержкой списков."""
        selected = self.parent.chat_widget.get_selected_text()
        if not selected:
            messagebox.showwarning("Сохранение", "Ничего не выделено в чате.")
            return
        first_line = selected.splitlines()[0].strip() if selected else ""
        if not first_line:
            first_line = "selected_text"
        base_name = self._safe_filename(first_line, "selected_text")
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Документ Word", "*.docx"), ("Все файлы", "*.*")],
            initialfile=base_name + ".docx",
            title="Сохранить выделенный текст как DOCX"
        )
        if not file_path:
            return
        doc = Document()
        lines = selected.splitlines()
        list_stack = []

        def close_lists_up_to(level):
            nonlocal list_stack
            while list_stack and list_stack[-1] >= level:
                list_stack.pop()

        def add_list_item(level, content):
            nonlocal list_stack
            close_lists_up_to(level)
            p = doc.add_paragraph()
            if level == 0:
                p.style = 'List Bullet'
            elif level == 1:
                p.style = 'List Bullet 2'
            elif level == 2:
                p.style = 'List Bullet 3'
            else:
                p.style = 'List Bullet 4'
            p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
            self._apply_inline_formatting(p, content)
            list_stack.append(level)

        def add_regular_paragraph(text):
            nonlocal list_stack
            list_stack = []
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            self._apply_inline_formatting(p, text)

        in_code_block = False
        code_block_lines = []
        i = 0
        while i < len(lines):
            line = lines[i]
            i += 1
            if line.strip().startswith('```'):
                if in_code_block:
                    if code_block_lines:
                        p = doc.add_paragraph()
                        run = p.add_run('\n'.join(code_block_lines))
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        code_block_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            if in_code_block:
                code_block_lines.append(line)
                continue
            if not line.strip():
                doc.add_paragraph()
                continue
            stripped = line.lstrip()
            indent = len(line) - len(stripped)
            level = indent // 4
            is_bullet = stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('• ')
            if is_bullet:
                content = stripped[2:].strip()
                add_list_item(level, content)
            else:
                header_match = re.match(r'^(#+)\s+(.*)$', line)
                if header_match:
                    list_stack = []
                    level = len(header_match.group(1))
                    text = header_match.group(2)
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.bold = True
                    if level == 1:
                        run.font.size = Pt(18)
                        p.style = 'Heading 1'
                    elif level == 2:
                        run.font.size = Pt(16)
                        p.style = 'Heading 2'
                    elif level == 3:
                        run.font.size = Pt(14)
                        p.style = 'Heading 3'
                    elif level >= 4:
                        run.font.size = Pt(12)
                        p.style = 'Heading 4'
                else:
                    number_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
                    if number_match:
                        content = number_match.group(2)
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(0)
                        run = p.add_run(f"{number_match.group(1)}. ")
                        run.bold = True
                        self._apply_inline_formatting(p, content)
                    else:
                        add_regular_paragraph(line)
        list_stack = []
        try:
            doc.save(file_path)
            self.parent.statusbar.config(text=f"Выделенный текст сохранён в DOCX: {file_path}")
            self.parent.add_system_message(f"📄 Выделенный текст сохранён в DOCX: {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка сохранения", f"Не удалось сохранить файл:\n{e}")