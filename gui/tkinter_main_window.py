# -*- coding: utf-8 -*-
"""Главное окно на Tkinter + ttkbootstrap (с PanedWindow, контекстным меню, Ctrl+C/V)"""

import logging
import threading
from pathlib import Path
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, Frame, Label, Text, Toplevel, ttk, Menu
from tkinter import END, DISABLED, NORMAL, LEFT, RIGHT, X, Y, BOTH, HORIZONTAL, VERTICAL
from tkinter import filedialog, messagebox, scrolledtext, Frame, Label, Text, Toplevel, ttk, Menu, simpledialog, scrolledtext

import ttkbootstrap as tb
from ttkbootstrap.constants import *
import markdown

from config.config_manager import ConfigManager
from core.chat_client import ChatGPTClient
from gui.file_loader import FileLoader
from gui.dialog_manager import DialogManager

from typing import Optional

import pandas as pd
from openpyxl import Workbook
import io
import json

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

# Импорт оркестратора и редактора сценариев
from orchestrator.dispatcher import run_analysis
from gui.scenario_editor import launch_scenario_editor

# Добавить для поддержки команд
import os
import glob
from typing import Optional


class ChatWidget(Frame):
    """Виджет чата с поддержкой Ctrl+C/X/V на любой раскладке"""

    def __init__(self, parent, on_send_callback):
        super().__init__(parent)
        self.parent = parent
        self.on_send_callback = on_send_callback

        self.paned = ttk.PanedWindow(self, orient=VERTICAL)
        self.paned.pack(fill=BOTH, expand=True)

        # === Область чата ===
        self.chat_area = scrolledtext.ScrolledText(
            self.paned, wrap=WORD, state=DISABLED,
            font=("Segoe UI", 10)
        )
        self.paned.add(self.chat_area, weight=3)

        # Контекстное меню и Ctrl+C для чата
        self._setup_chat_context_menu()
        self.chat_area.bind("<Control-KeyPress>", self._on_ctrl_key_chat)

        # === Панель ввода ===
        input_frame = Frame(self.paned)
        self.paned.add(input_frame, weight=1)

        self.input_field = Text(input_frame, height=4, wrap=WORD, font=("Segoe UI", 10))
        self.input_field.pack(side=LEFT, fill=BOTH, expand=True, padx=(0, 5))

        # Привязка Ctrl+KeyPress для поля ввода
        self.input_field.bind("<Control-KeyPress>", self._on_ctrl_key_input)

        self.send_btn = tb.Button(
            input_frame, text="Отправить", bootstyle="primary",
            command=self._send_message
        )
        self.send_btn.pack(side=RIGHT)

        self.input_field.bind("<Control-Return>", lambda e: self._send_message())

        # --- Кнопки загрузки/сохранения промпта ---
        self.load_prompt_btn = tb.Button(
            input_frame, text="📂", bootstyle="secondary",
            command=self._load_prompt_from_dialog, width=3
        )
        self.load_prompt_btn.pack(side=RIGHT, padx=1)

        self.save_prompt_btn = tb.Button(
            input_frame, text="💾", bootstyle="secondary",
            command=self._save_prompt_to_dialog, width=3
        )
        self.save_prompt_btn.pack(side=RIGHT, padx=1)

        # Теги форматирования
        self.chat_area.tag_config("user_header", font=("Segoe UI", 9, "bold"), foreground="blue")
        self.chat_area.tag_config("assistant_header", font=("Segoe UI", 9, "bold"), foreground="green")
        self.chat_area.tag_config("user_text", font=("Segoe UI", 10), lmargin1=10, lmargin2=10, rmargin=10)
        self.chat_area.tag_config("assistant_text", font=("Segoe UI", 10), lmargin1=10, lmargin2=10, rmargin=10)

    # ---------- Контекстное меню чата ----------
    def _setup_chat_context_menu(self):
        self.chat_context_menu = Menu(self.chat_area, tearoff=0)
        self.chat_context_menu.add_command(label="Копировать выделенное", command=self._copy_selected_from_chat)
        self.chat_context_menu.add_command(label="Копировать всё", command=self._copy_all_from_chat)
        self.chat_area.bind("<Button-3>", self._show_chat_context_menu)

    def _show_chat_context_menu(self, event):
        try:
            self.chat_context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.chat_context_menu.grab_release()

    def _copy_selected_from_chat(self):
        try:
            self.chat_area.configure(state=NORMAL)
            selected = self.chat_area.get("sel.first", "sel.last")
            self.chat_area.configure(state=DISABLED)
            if selected:
                self.clipboard_clear()
                self.clipboard_append(selected)
        except:
            pass

    def _copy_all_from_chat(self):
        self.chat_area.configure(state=NORMAL)
        all_text = self.chat_area.get("1.0", END)
        self.chat_area.configure(state=DISABLED)
        if all_text.strip():
            self.clipboard_clear()
            self.clipboard_append(all_text)

    # ---------- Обработчики Ctrl+Key (все раскладки) ----------
    def _on_ctrl_key_chat(self, event):
        """Ctrl+клавиша для чата"""
        if event.keycode == 67:  # C
            self._copy_from_chat()
            return "break"
        return None

    def _on_ctrl_key_input(self, event):
        """Ctrl+клавиша для поля ввода"""
        if event.keycode == 67:  # C – копировать
            self._copy_from_input()
            return "break"
        elif event.keycode == 88:  # X – вырезать
            self._cut_from_input()
            return "break"
        elif event.keycode == 86:  # V – вставить
            self._paste_to_input()
            return "break"
        return None

    # ---------- Функции для чата ----------
    def _copy_from_chat(self):
        try:
            self.chat_area.configure(state=NORMAL)
            selected = self.chat_area.get("sel.first", "sel.last")
            self.chat_area.configure(state=DISABLED)
            if selected:
                self.clipboard_clear()
                self.clipboard_append(selected)
        except:
            pass

    # ---------- Функции для поля ввода ----------
    def _copy_from_input(self):
        """Копировать выделенное из поля ввода"""
        try:
            # Пробуем стандартное событие
            self.input_field.event_generate("<<Copy>>")
        except:
            # Ручное копирование
            try:
                selected = self.input_field.get("sel.first", "sel.last")
                if selected:
                    self.clipboard_clear()
                    self.clipboard_append(selected)
            except:
                pass

    def _cut_from_input(self):
        """Вырезать выделенное из поля ввода"""
        try:
            self._copy_from_input()
            self.input_field.delete("sel.first", "sel.last")
        except:
            pass

    def _paste_to_input(self):
        """Вставить из буфера в поле ввода"""
        try:
            self.input_field.event_generate("<<Paste>>")
        except:
            try:
                text = self.clipboard_get()
                self.input_field.insert("insert", text)
            except:
                pass

    # ---------- Основные методы ----------
    def _send_message(self):
        text = self.input_field.get("1.0", END).strip()
        if text:
            self.on_send_callback(text)
            self.input_field.delete("1.0", END)

    def add_message(self, text: str, sender: str):
        self.chat_area.configure(state=NORMAL)
        timestamp = datetime.now().strftime("%H:%M")
        header = "Вы" if sender == "user" else "Ассистент"
        header_tag = "user_header" if sender == "user" else "assistant_header"
        text_tag = "user_text" if sender == "user" else "assistant_text"
        self.chat_area.insert(END, f"{header}  {timestamp}\n", header_tag)
        self.chat_area.insert(END, f"{text}\n", text_tag)
        self.chat_area.insert(END, "-" * 60 + "\n", "separator")
        self.chat_area.tag_config("separator", foreground="gray", font=("Segoe UI", 8))
        self.chat_area.see(END)
        self.chat_area.configure(state=DISABLED)

    def get_chat_text(self) -> str:
        self.chat_area.configure(state=NORMAL)
        text = self.chat_area.get("1.0", END)
        self.chat_area.configure(state=DISABLED)
        return text.strip()

    def clear(self):
        self.chat_area.configure(state=NORMAL)
        self.chat_area.delete("1.0", END)
        self.chat_area.configure(state=DISABLED)

    def set_send_enabled(self, enabled: bool):
        self.send_btn.configure(state=NORMAL if enabled else DISABLED)

    def _load_prompt_from_dialog(self):
        file_path = filedialog.askopenfilename(
            title="Загрузить промпт",
            filetypes=[("Текстовые файлы", "*.txt *.md *.prompt"), ("Все файлы", "*.*")]
        )
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.input_field.delete("1.0", END)
                self.input_field.insert("1.0", content)
                self.on_send_callback(f"*Система*: Промпт загружен из {file_path}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось загрузить промпт:\n{e}")

    def _save_prompt_to_dialog(self):
        text = self.input_field.get("1.0", END).strip()
        if not text:
            messagebox.showwarning("Сохранение", "Поле ввода пустое")
            return
        file_path = filedialog.asksaveasfilename(
            defaultextension=".prompt",
            filetypes=[("Текстовые файлы", "*.txt *.md *.prompt"), ("Все файлы", "*.*")],
            title="Сохранить промпт"
        )
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                self.on_send_callback(f"*Система*: Промпт сохранён в {file_path}")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось сохранить промпт:\n{e}")

class SettingsDialog(Toplevel):
    """Диалог настроек с выпадающим списком моделей и настройками диалога"""

    def __init__(self, parent, config_manager: ConfigManager, available_models: list):
        super().__init__(parent)
        self.config_manager = config_manager
        self.available_models = available_models
        self.title("Настройки")
        self.geometry("500x550")
        self.resizable(False, False)
        self.transient(parent)
        self.grab_set()

        # Основной фрейм с прокруткой
        main_frame = Frame(self)
        main_frame.pack(fill=BOTH, expand=True, padx=10, pady=10)

        # Контейнер с прокруткой
        canvas = tk.Canvas(main_frame)
        scrollbar = ttk.Scrollbar(main_frame, orient=VERTICAL, command=canvas.yview)
        self.scrollable_frame = Frame(canvas)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side=LEFT, fill=BOTH, expand=True)
        scrollbar.pack(side=RIGHT, fill=Y)

        # --- Модель (выпадающий список) ---
        tb.Label(self.scrollable_frame, text="Модель:").pack(anchor=W, pady=(0, 5))
        self.model_combo = tb.Combobox(self.scrollable_frame, state="readonly", width=50)
        self.model_combo.pack(fill=X, pady=(0, 10))

        # Заполняем список моделей
        models_to_show = self.available_models if self.available_models else ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo", "gpt-4"]
        self.model_combo['values'] = models_to_show
        current_model = self.config_manager.config.get("llm_defaults", {}).get("model", "gpt-4o-mini")
        if current_model in models_to_show:
            self.model_combo.set(current_model)
        else:
            self.model_combo.set(models_to_show[0] if models_to_show else "gpt-4o-mini")

        # --- Температура ---
        tb.Label(self.scrollable_frame, text="Температура (0.0-2.0):").pack(anchor=W, pady=(0, 5))
        temp_frame = Frame(self.scrollable_frame)
        temp_frame.pack(fill=X, pady=(0, 10))
        self.temp_scale = tb.Scale(
            temp_frame,
            from_=0.0,
            to=2.0,
            value=self.config_manager.config.get("llm_defaults", {}).get("temperature", 0.3),
            orient=HORIZONTAL
        )
        self.temp_scale.pack(side=LEFT, fill=X, expand=True)
        self.temp_label = tb.Label(temp_frame, text=f"{self.temp_scale.get():.1f}", width=5)
        self.temp_label.pack(side=RIGHT, padx=(5, 0))
        self.temp_scale.configure(command=self._update_temp_label)

        # --- Max tokens ---
        tb.Label(self.scrollable_frame, text="Max tokens:").pack(anchor=W, pady=(0, 5))
        self.tokens_spin = tb.Spinbox(self.scrollable_frame, from_=100, to=8000, width=15)
        self.tokens_spin.set(str(self.config_manager.config.get("llm_defaults", {}).get("max_tokens", 2000)))
        self.tokens_spin.pack(anchor=W, pady=(0, 10))

        # --- Настройки диалога ---
        dialog_params = self.config_manager.get_dialog_params()
        tb.Label(self.scrollable_frame, text="Настройки диалога:", font=("Segoe UI", 10, "bold")).pack(anchor=W, pady=(10, 5))

        frame_dialog = Frame(self.scrollable_frame)
        frame_dialog.pack(fill=X, pady=5)

        tb.Label(frame_dialog, text="Пар для контекста:").grid(row=0, column=0, sticky=W, padx=5, pady=2)
        self.history_pairs_spin = tb.Spinbox(frame_dialog, from_=1, to=20, width=5)
        self.history_pairs_spin.set(dialog_params.get("history_pairs", 3))
        self.history_pairs_spin.grid(row=0, column=1, padx=5, pady=2)

        tb.Label(frame_dialog, text="Пар для summary:").grid(row=1, column=0, sticky=W, padx=5, pady=2)
        self.summary_pairs_spin = tb.Spinbox(frame_dialog, from_=1, to=20, width=5)
        self.summary_pairs_spin.set(dialog_params.get("summary_pairs", 3))
        self.summary_pairs_spin.grid(row=1, column=1, padx=5, pady=2)

        tb.Label(frame_dialog, text="Модель для summary:").grid(row=2, column=0, sticky=W, padx=5, pady=2)
        self.summary_model_combo = tb.Combobox(frame_dialog, state="readonly", width=30)
        self.summary_model_combo['values'] = self.available_models if self.available_models else ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo", "gpt-4"]
        current_summary_model = dialog_params.get("summary_model", "gpt-4o-mini")
        if current_summary_model in self.summary_model_combo['values']:
            self.summary_model_combo.set(current_summary_model)
        else:
            self.summary_model_combo.set(self.summary_model_combo['values'][0])
        self.summary_model_combo.grid(row=2, column=1, padx=5, pady=2)

        # --- Тема интерфейса ---
        tb.Label(self.scrollable_frame, text="Тема:", font=("Segoe UI", 10, "bold")).pack(anchor=W, pady=(10, 5))
        self.theme_combo = tb.Combobox(self.scrollable_frame, values=["light", "dark"], state="readonly", width=20)
        self.theme_combo.set(self.config_manager.config.get("ui", {}).get("theme", "light"))
        self.theme_combo.pack(anchor=W, pady=(0, 10))

        # --- Кнопки ---
        btn_frame = Frame(self.scrollable_frame)
        btn_frame.pack(fill=X, pady=(20, 10))
        tb.Button(btn_frame, text="Сохранить", bootstyle="primary", command=self._save).pack(side=RIGHT, padx=5)
        tb.Button(btn_frame, text="Отмена", bootstyle="secondary", command=self.destroy).pack(side=RIGHT, padx=5)

    def _update_temp_label(self, *args):
        self.temp_label.config(text=f"{self.temp_scale.get():.1f}")

    def _save(self):
        changes = {
            "llm_defaults": {
                "model": self.model_combo.get(),
                "temperature": self.temp_scale.get(),
                "max_tokens": int(self.tokens_spin.get())
            },
            "dialog": {
                "history_pairs": int(self.history_pairs_spin.get()),
                "summary_pairs": int(self.summary_pairs_spin.get()),
                "summary_model": self.summary_model_combo.get().strip()
            },
            "ui": {
                "theme": self.theme_combo.get()
            }
        }
        self.config_manager.update_config(changes)
        self.destroy()

class MainWindow(tb.Window):
    """Главное окно приложения"""

    def __init__(self, config_manager: ConfigManager):
        # Выбор темы из конфига
        theme = config_manager.config.get("ui", {}).get("theme", "light")
        theme_name = "darkly" if theme == "dark" else "cosmo"

        super().__init__(themename=theme_name)
        self.config_manager = config_manager
        self.chat_client = ChatGPTClient(config_manager)
        self.dialog_manager = DialogManager(config_manager, self.chat_client)
        self._stream_buffer = ""
        self.logger = logging.getLogger("main_window")
        self.analysis_thread = None
        self.cancel_event = None
        self.result_format = "text"
        self.setup_logging()
        self.init_ui()
        self.bind_signals()
        self.check_api_key()
        self.after(500, self.test_api_connection)
        self.after(1000, self._auto_load_prompt)

        self.available_models = []  # список моделей, загруженных при старте
        self.load_models_at_startup()

        self.attached_files = []  # список кортежей (filename, content) для текущего сеанса

    def setup_logging(self):
        logs_dir = self.config_manager.config.get("paths", {}).get("logs_dir", "./logs")
        Path(logs_dir).mkdir(exist_ok=True)
        log_file = Path(logs_dir) / "app.log"
        logging.basicConfig(
            level=logging.INFO,
            format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
            handlers=[
                logging.FileHandler(log_file, encoding="utf-8"),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger("main_window")
        # Ниже добавляем строку:
        logging.getLogger("chat_client").setLevel(logging.DEBUG)

    def init_ui(self):
        self.title("AI Document Analyst (Chat)")
        self.geometry("900x600")
        self.minsize(600, 400)

        # Меню
        menubar = tb.Menu(self)
        self.config(menu=menubar)

        # Файл
        file_menu = tb.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Файл", menu=file_menu)
        file_menu.add_command(label="Выход", command=self.destroy)

        # Диалог
        dialog_menu = tb.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Диалог", menu=dialog_menu)
        dialog_menu.add_command(label="Новый диалог", command=self.new_dialog)
        dialog_menu.add_command(label="Сохранить диалог", command=self.save_dialog)
        dialog_menu.add_command(label="Загрузить диалог", command=self.load_dialog)
        dialog_menu.add_command(label="Обобщить диалог (summary)", command=self.summarize_dialog)
        dialog_menu.add_separator()
        dialog_menu.add_command(label="Очистить загруженные документы", command=self._cmd_clear_documents)
        dialog_menu.add_command(label="Очистить загруженные сценарии", command=self._cmd_clear_scenarios)
        dialog_menu.add_separator()
        dialog_menu.add_command(label="Загрузить промпт", command=self._load_prompt_global)
        dialog_menu.add_command(label="Сохранить промпт", command=self._save_prompt_global)

        # Инструменты
        tools_menu = tb.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Инструменты", menu=tools_menu)
        tools_menu.add_command(label="Копировать чат", command=self.copy_chat)
        tools_menu.add_command(label="Сохранить чат в DOCX", command=self._save_chat_as_docx_wrapper)
        tools_menu.add_separator()
        tools_menu.add_command(label="Очистить чат", command=self.clear_chat)
        tools_menu.add_command(label="Настройки", command=self.open_settings)

        # Команды
        commands_menu = tb.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Команды", menu=commands_menu)
        commands_menu.add_command(label="📄 Загрузить документ", command=self._cmd_load_document_dialog)
        commands_menu.add_command(label="📋 Загрузить сценарий", command=self._cmd_load_scenario_dialog)
        commands_menu.add_separator()
        commands_menu.add_command(label="▶️ Запустить сценарий", command=self._cmd_run_scenario)
        commands_menu.add_separator()
        commands_menu.add_command(label="📁 Импортировать папку", command=self._cmd_import_folder_dialog)
        commands_menu.add_separator()
        # Подменю "Экспортировать результат"
        export_menu = tb.Menu(commands_menu, tearoff=0)
        commands_menu.add_cascade(label="💾 Экспортировать результат", menu=export_menu)
        export_menu.add_command(label="JSON",command=lambda: self._cmd_export_result("json"))
        export_menu.add_command(label="CSV",command=lambda: self._cmd_export_result("csv"))
        export_menu.add_command(label="Markdown (MD)",command=lambda: self._cmd_export_result("md"))
        export_menu.add_command(label="Word (DOCX)",command=lambda: self._cmd_export_result("docx"))
        export_menu.add_command(label="Excel (XLSX)",command=lambda: self._cmd_export_result("xlsx"))
        commands_menu.add_separator()
        commands_menu.add_command(label="🗑️ Очистить документы", command=self._cmd_clear_documents)
        commands_menu.add_command(label="🗑️ Очистить сценарии", command=self._cmd_clear_scenarios)
        commands_menu.add_separator()
        commands_menu.add_command(label="❓ Помощь", command=self._cmd_help)

        help_menu = tb.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Помощь", menu=help_menu)
        help_menu.add_command(label="О программе", command=self.show_about)

        # Чат
        # self.chat_widget = ChatWidget(self, self.on_user_message)
        # self.chat_widget.pack(fill=BOTH, expand=True)

        # +++ Notebook с вкладками
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill=BOTH, expand=True)

        # Вкладка "Чат"
        self.chat_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.chat_frame, text="Чат")
        self.chat_widget = ChatWidget(self.chat_frame, self.on_user_message)
        self.chat_widget.pack(fill=BOTH, expand=True)

        # Вкладка "Анализ документа"
        self.analysis_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.analysis_frame, text="Анализ документа")
        self._build_analysis_tab()  # новый метод

        # Панель загрузки файла и дополнительных кнопок
        file_frame = Frame(self)
        file_frame.pack(fill=X, padx=5, pady=5)

        self.attach_btn = tb.Button(
            file_frame, text="📎 Прикрепить файл", bootstyle="secondary",
            command=self.load_file
        )
        self.attach_btn.pack(side=LEFT, padx=2)

        self.file_label = tb.Label(file_frame, text="Файл не загружен", bootstyle="info")
        self.file_label.pack(side=LEFT, padx=10)

        self.clear_file_btn = tb.Button(
            file_frame, text="Убрать файл", bootstyle="danger",
            command=self.clear_file, state=DISABLED
        )
        self.clear_file_btn.pack(side=LEFT, padx=2)

        # Кнопка копирования чата
        self.copy_chat_btn = tb.Button(
            file_frame, text="📋 Копировать чат", bootstyle="info",
            command=self.copy_chat
        )
        self.copy_chat_btn.pack(side=RIGHT, padx=2)

        # Кнопка сохранения чата в DOCX
        self.save_chat_docx_btn = tb.Button(
            file_frame, text="💾 Сохранить чат (DOCX)", bootstyle="info",
            command=self._save_chat_as_docx_wrapper
        )
        self.save_chat_docx_btn.pack(side=RIGHT, padx=2)

        # Статусная строка
        self.statusbar = tb.Label(self, text="Готов", bootstyle="info", anchor=W)
        self.statusbar.pack(side=BOTTOM, fill=X)

        # Приветственное сообщение
        self.chat_widget.add_message("Приложение запущено. Введите вопрос.", "assistant")

    def _build_analysis_tab(self):
        """Строит интерфейс для анализа документов по сценарию"""
        # Панель выбора документа
        doc_frame = ttk.LabelFrame(self.analysis_frame, text="Документ", padding=5)
        doc_frame.pack(fill=X, padx=5, pady=5)
        self.doc_path_var = tb.StringVar()
        self.doc_entry = tb.Entry(doc_frame, textvariable=self.doc_path_var, state="readonly")
        self.doc_entry.pack(side=LEFT, fill=X, expand=True, padx=(0, 5))
        tb.Button(doc_frame, text="Выбрать файл", bootstyle="secondary",
                  command=self._select_document).pack(side=RIGHT)

        # Панель выбора сценария
        scenario_frame = ttk.LabelFrame(self.analysis_frame, text="Сценарий", padding=5)
        scenario_frame.pack(fill=X, padx=5, pady=5)
        self.scenario_path_var = tb.StringVar()
        self.scenario_entry = tb.Entry(scenario_frame, textvariable=self.scenario_path_var, state="readonly")
        self.scenario_entry.pack(side=LEFT, fill=X, expand=True, padx=(0, 5))
        btn_scenario = tb.Button(scenario_frame, text="Загрузить JSON", bootstyle="secondary",
                                 command=self._select_scenario)
        btn_scenario.pack(side=RIGHT, padx=2)
        tb.Button(scenario_frame, text="Редактор сценариев", bootstyle="info",
                  command=self._open_scenario_editor).pack(side=RIGHT, padx=2)

        # Кнопки управления
        control_frame = ttk.Frame(self.analysis_frame)
        control_frame.pack(fill=X, padx=5, pady=5)
        self.analyze_btn = tb.Button(control_frame, text="Запустить анализ", bootstyle="success",
                                     command=self._start_analysis, state=DISABLED)
        self.analyze_btn.pack(side=LEFT, padx=2)
        self.cancel_btn = tb.Button(control_frame, text="Отменить", bootstyle="danger",
                                    command=self._cancel_analysis, state=DISABLED)
        self.cancel_btn.pack(side=LEFT, padx=2)

        # Прогресс
        self.progress_var = tb.IntVar(value=0)
        self.progress_bar = ttk.Progressbar(self.analysis_frame, variable=self.progress_var,
                                            mode='determinate')
        self.progress_bar.pack(fill=X, padx=5, pady=5)

        # --- Вкладки для вывода информации ---
        output_notebook = ttk.Notebook(self.analysis_frame)
        output_notebook.pack(fill=BOTH, expand=True, padx=5, pady=5)

        # Вкладка "Лог выполнения"
        log_frame = ttk.Frame(output_notebook)
        output_notebook.add(log_frame, text="Лог выполнения")
        self.log_text = scrolledtext.ScrolledText(log_frame, height=8, wrap=WORD, state=DISABLED)
        self.log_text.pack(fill=BOTH, expand=True)

        # Вкладка "Результат анализа"
        result_frame = ttk.Frame(output_notebook)
        output_notebook.add(result_frame, text="Результат анализа")
        self.result_text = scrolledtext.ScrolledText(result_frame, height=10, wrap=WORD, state=DISABLED)
        self.result_text.pack(fill=BOTH, expand=True)

        # Панель кнопок для результата (размещаем внутри вкладки, можно добавить)
        btn_result_frame = ttk.Frame(result_frame)
        btn_result_frame.pack(fill=X, pady=5)
        tb.Button(btn_result_frame, text="Копировать результат", bootstyle="secondary",
                  command=self._copy_result).pack(side=LEFT, padx=2)
        tb.Button(btn_result_frame, text="Сохранить результат", bootstyle="secondary",
                  command=self._save_result).pack(side=LEFT, padx=2)

        # Вкладка "Промежуточные результаты этапов"
        intermediate_frame = ttk.Frame(output_notebook)
        output_notebook.add(intermediate_frame, text="Промежуточные результаты")
        self.intermediate_combo = ttk.Combobox(intermediate_frame, state="readonly", width=30)
        self.intermediate_combo.pack(fill=X, pady=(5, 5))
        self.intermediate_combo.bind("<<ComboboxSelected>>", self._on_intermediate_selected)
        self.intermediate_text = scrolledtext.ScrolledText(intermediate_frame, height=8, wrap=WORD,
                                                           state=DISABLED)
        self.intermediate_text.pack(fill=BOTH, expand=True)

        # Переменные для управления потоком
        self.analysis_thread = None
        self.cancel_event = None
    def bind_signals(self):
        """Устанавливаем колбэки для ChatGPTClient"""
        self.chat_client.on_response = self._on_ai_response
        self.chat_client.on_chunk = self._on_ai_chunk
        self.chat_client.on_error = self._on_api_error
        self.chat_client.on_start = self._on_request_start
        self.chat_client.on_finish = self._on_request_finish

    def _select_document(self):
        file_path = filedialog.askopenfilename(
            title="Выберите документ",
            filetypes=[
                ("Поддерживаемые", "*.docx *.xlsx *.pdf *.xls *.txt *.md"),
                ("Документы Word", "*.docx"),
                ("Таблицы Excel", "*.xlsx *.xls"),
                ("PDF", "*.pdf"),
                ("Текстовые", "*.txt *.md"),
                ("Все файлы", "*.*")
            ]
        )
        if file_path:
            self.doc_path_var.set(file_path)
            self._update_analyze_button_state()

    def _select_scenario(self):
        file_path = filedialog.askopenfilename(
            title="Выберите JSON-сценарий",
            filetypes=[("JSON files", "*.json")]
        )
        if file_path:
            self.scenario_path_var.set(file_path)
            # Простая валидация
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    if 'stages' not in data:
                        raise ValueError("Нет ключа 'stages'")
                self._update_analyze_button_state()
            except Exception as e:
                messagebox.showerror("Ошибка", f"Невалидный сценарий:\n{e}")
                self.scenario_path_var.set("")

    def _update_analyze_button_state(self):
        enabled = bool(self.doc_path_var.get() and self.scenario_path_var.get())
        self.analyze_btn.config(state=NORMAL if enabled else DISABLED)

    def _open_scenario_editor(self):
        launch_scenario_editor(self, self.config_manager)

    def _start_analysis(self):
        if self.analysis_thread and self.analysis_thread.is_alive():
            messagebox.showwarning("Анализ", "Анализ уже выполняется")
            return

        doc_path = self.doc_path_var.get()
        scenario_path = self.scenario_path_var.get()
        if not doc_path or not scenario_path:
            return

        # Блокируем кнопки
        self.analyze_btn.config(state=DISABLED)
        self.cancel_btn.config(state=NORMAL)
        self.progress_var.set(0)
        self.log_text.config(state=NORMAL)
        self.log_text.delete("1.0", END)
        self.log_text.insert(END, "Начало анализа...\n")
        self.log_text.config(state=DISABLED)
        self.result_text.config(state=NORMAL)
        self.result_text.delete("1.0", END)
        self.result_text.config(state=DISABLED)

        # Очищаем промежуточные результаты
        self._intermediate_results = {}
        self.intermediate_combo['values'] = []
        self.intermediate_combo.set("")
        self.intermediate_text.config(state=NORMAL)
        self.intermediate_text.delete("1.0", END)
        self.intermediate_text.config(state=DISABLED)

        # Получаем настройки API
        api_key = self.config_manager.get_api_key()
        if not api_key:
            self._append_log("❌ API ключ не найден. Задайте его в настройках.")
            self._end_analysis(False)
            return

        llm_params = self.config_manager.get_llm_params()
        api_settings = {
            "api_key": api_key,
            "model": llm_params.get("model", "gpt-4o-mini"),
            "temperature": llm_params.get("temperature", 0.2),
            "max_tokens": llm_params.get("max_tokens", 2000),
            "timeout": llm_params.get("timeout", 60)
        }

        self.cancel_event = threading.Event()
        self.analysis_thread = threading.Thread(
            target=self._run_analysis_thread,
            args=(doc_path, scenario_path, api_settings),
            daemon=True
        )
        self.analysis_thread.start()

    def _run_analysis_thread(self, doc_path, scenario_path, api_settings):
        """Выполняется в отдельном потоке, вызывает оркестратор"""

        def progress_callback(msg_type, data):
            # msg_type может быть "log" или "stage"
            self.after(0, lambda: self._handle_progress(msg_type, data))

        try:
            result = run_analysis(
                document_path=doc_path,
                scenario_path=scenario_path,
                api_settings=api_settings,
                progress_callback=progress_callback,
                cancel_event=self.cancel_event
            )
            self.after(0, lambda: self._on_analysis_finished(result))
        except Exception as e:
            self.after(0, lambda: self._on_analysis_error(str(e)))

    def _handle_progress(self, msg_type, data):
        print(f"DEBUG: msg_type={msg_type}, data={data}")
        if msg_type == "log":
            self._append_log(data)
        elif msg_type == "stage":
            self._update_progress(data["current"], data["total"], data["name"])
        elif msg_type == "stage_result":
            stage_id = data.get("stage_id", "unknown")
            preview = data.get("result_preview", "")
            self._append_log(f"--- Результат этапа '{stage_id}': ---")
            self._append_log(preview)
            self._append_log("---")
            # Сохраняем полный результат для возможного просмотра позже
            if not hasattr(self, '_intermediate_results'):
                self._intermediate_results = {}
            self._intermediate_results[stage_id] = data.get("full_result", preview)
            self.after(0, self._update_intermediate_combo)

    def _update_progress(self, current, total, stage_name):
        percent = int((current / total) * 100)
        self.progress_var.set(percent)
        self._append_log(f"Этап {current}/{total}: {stage_name}")

    def _append_log(self, message):
        self.log_text.config(state=NORMAL)
        self.log_text.insert(END, f"{message}\n")
        self.log_text.see(END)
        self.log_text.config(state=DISABLED)

    def _on_analysis_finished(self, result):
        self._end_analysis(True)
        if result["status"] == "success":
            self._append_log("✅ Анализ завершён успешно")
            self._show_result(result["result"], result.get("format", "text"))
        elif result["status"] == "cancelled":
            self._append_log("⚠️ Анализ отменён пользователем")
        else:
            self._append_log(f"❌ Ошибка: {result.get('message', 'Неизвестная ошибка')}")

    def _on_analysis_error(self, error_msg):
        self._end_analysis(False)
        self._append_log(f"❌ Критическая ошибка: {error_msg}")

    def _end_analysis(self, success):
        self.after(0, lambda: self.analyze_btn.config(
            state=NORMAL if self.doc_path_var.get() and self.scenario_path_var.get() else DISABLED))
        self.after(0, lambda: self.cancel_btn.config(state=DISABLED))
        self.analysis_thread = None
        self.cancel_event = None

    def _cancel_analysis(self):
        if self.cancel_event:
            self.cancel_event.set()
            self._append_log("Отмена анализа...")

    def _show_result(self, text, fmt):
        self.result_text.config(state=NORMAL)
        self.result_text.delete("1.0", END)
        # Для markdown можно сделать рендеринг (упрощённо, как обычный текст)
        # Или оставить как есть
        self.result_text.insert(END, text)
        self.result_text.config(state=DISABLED)
        self.result_format = fmt

    def _copy_result(self):
        text = self.result_text.get("1.0", END).strip()
        if text:
            self.clipboard_clear()
            self.clipboard_append(text)
            self.statusbar.config(text="Результат скопирован в буфер обмена")

    def _save_result(self):
        text = self.result_text.get("1.0", END).strip()
        if not text:
            return
        fmt = getattr(self, 'result_format', 'text')

        filetypes = [
            ("Текстовый файл", "*.txt"),
            ("Markdown", "*.md"),
            ("JSON", "*.json"),
            ("CSV", "*.csv"),
            ("Документ Word", "*.docx"),
            ("Таблица Excel", "*.xlsx")
        ]
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=filetypes
        )
        if not file_path:
            return

        ext = Path(file_path).suffix.lower()
        try:
            if ext == '.docx':
                self._save_as_docx(file_path, text, fmt)
            elif ext == '.xlsx':
                self._save_as_xlsx(file_path, text, fmt)
            else:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                self.statusbar.config(text=f"Результат сохранён: {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка сохранения", f"Не удалось сохранить файл:\n{e}")
            self.logger.error(f"Ошибка сохранения: {e}")

    def _parse_markdown_table(self, table_lines):
        """Парсит строки markdown-таблицы в список списков."""
        lines = [line.strip() for line in table_lines if line.strip()]
        if len(lines) < 2:
            return None
        header_line = lines[0]
        sep_line = lines[1] if len(lines) > 1 else None
        if not sep_line or '---' not in sep_line:
            data_rows = lines
            headers = []
        else:
            headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
            data_rows = lines[2:]
        table = []
        if headers:
            table.append(headers)
        for row_line in data_rows:
            cells = [cell.strip() for cell in row_line.split('|') if cell.strip()]
            if len(cells) < len(headers):
                cells.extend([''] * (len(headers) - len(cells)))
            table.append(cells[:len(headers)])
        return table

    def _save_as_docx(self, file_path: str, content: str, fmt: str):
        """Сохраняет результат в формате DOCX."""
        doc = Document()

        if fmt == 'markdown':
            lines = content.split('\n')
            table_lines = []
            in_table = False
            for line in lines:
                if '|' in line and not line.strip().startswith('|'):
                    if not in_table:
                        in_table = True
                    table_lines.append(line)
                else:
                    if in_table:
                        if table_lines:
                            table = self._parse_markdown_table(table_lines)
                            if table and len(table) > 0 and len(table[0]) > 0:
                                word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                                word_table.style = 'Table Grid'
                                for i, row_data in enumerate(table):
                                    for j, cell_text in enumerate(row_data):
                                        word_table.cell(i, j).text = cell_text
                        table_lines = []
                        in_table = False
                        if line.strip():
                            doc.add_paragraph(line)
                    else:
                        if line.strip():
                            doc.add_paragraph(line)
            if table_lines:
                table = self._parse_markdown_table(table_lines)
                if table and len(table) > 0 and len(table[0]) > 0:
                    word_table = doc.add_table(rows=len(table), cols=len(table[0]))
                    word_table.style = 'Table Grid'
                    for i, row_data in enumerate(table):
                        for j, cell_text in enumerate(row_data):
                            word_table.cell(i, j).text = cell_text

        elif fmt == 'json':
            try:
                data = json.loads(content)
                self._json_to_word_table(doc, data)
            except json.JSONDecodeError:
                # Если невалидный JSON, вставляем как текст
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

        else:
            # CSV, text – вставляем моноширинным шрифтом
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

        doc.save(file_path)
        self.statusbar.config(text=f"Результат сохранён в DOCX: {file_path}")

    def _save_as_xlsx(self, file_path: str, content: str, fmt: str):
        """Сохраняет результат в формате XLSX."""
        wb = Workbook()

        if fmt in ('json', 'csv'):
            try:
                if fmt == 'json':
                    data = json.loads(content)
                    if isinstance(data, list):
                        df = pd.DataFrame(data)
                    else:
                        df = pd.DataFrame([data])
                else:  # csv
                    df = pd.read_csv(io.StringIO(content))
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False, sheet_name='Data')
                self.statusbar.config(text=f"Результат сохранён в XLSX: {file_path}")
                return
            except Exception as e:
                self.logger.warning(f"Не удалось преобразовать в таблицу: {e}")
                # переходим к текстовому сохранению

        if fmt == 'markdown':
            lines = content.split('\n')
            table_lines = []
            in_table = False
            text_parts = []
            tables = []
            for line in lines:
                if '|' in line and not line.strip().startswith('|'):
                    if not in_table:
                        in_table = True
                    table_lines.append(line)
                else:
                    if in_table:
                        if table_lines:
                            table = self._parse_markdown_table(table_lines)
                            if table:
                                tables.append(table)
                        table_lines = []
                        in_table = False
                        text_parts.append(line)
                    else:
                        text_parts.append(line)
            if table_lines:
                table = self._parse_markdown_table(table_lines)
                if table:
                    tables.append(table)

            if text_parts:
                ws_text = wb.active
                ws_text.title = "Text"
                ws_text.cell(row=1, column=1, value='\n'.join(text_parts))

            if tables:
                ws_table = wb.create_sheet("Tables")
                start_row = 1
                for table in tables:
                    for row_data in table:
                        for col_idx, cell_value in enumerate(row_data, 1):
                            ws_table.cell(row=start_row, column=col_idx, value=cell_value)
                        start_row += 1
                    start_row += 1  # пустая строка между таблицами
        else:
            # text – просто вставляем в ячейку
            ws = wb.active
            ws.cell(row=1, column=1, value=content)

        wb.save(file_path)
        self.statusbar.config(text=f"Результат сохранён в XLSX: {file_path}")

    def _json_to_word_table(self, doc, data):
        """Преобразует JSON-данные в таблицу Word."""
        if isinstance(data, list) and data and all(isinstance(item, dict) for item in data):
            # Массив объектов → таблица с заголовками
            headers = list(data[0].keys())
            table = doc.add_table(rows=1 + len(data), cols=len(headers))
            table.style = 'Table Grid'
            # Заголовки
            for col_idx, header in enumerate(headers):
                table.cell(0, col_idx).text = header
            # Данные
            for row_idx, item in enumerate(data, 1):
                for col_idx, key in enumerate(headers):
                    val = item.get(key, '')
                    if isinstance(val, (list, dict)):
                        val = json.dumps(val, ensure_ascii=False)
                    table.cell(row_idx, col_idx).text = str(val)
        elif isinstance(data, dict):
            # Одиночный объект → таблица "ключ-значение"
            table = doc.add_table(rows=len(data), cols=2)
            table.style = 'Table Grid'
            for row_idx, (key, val) in enumerate(data.items()):
                table.cell(row_idx, 0).text = str(key)
                if isinstance(val, (list, dict)):
                    val = json.dumps(val, ensure_ascii=False)
                table.cell(row_idx, 1).text = str(val)
        else:
            # Другие типы → вставляем как текст
            p = doc.add_paragraph()
            run = p.add_run(json.dumps(data, ensure_ascii=False))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

    def _on_intermediate_selected(self, event):
        """Показывает полный результат выбранного этапа"""
        stage_id = self.intermediate_combo.get()
        if hasattr(self, '_intermediate_results') and stage_id in self._intermediate_results:
            self.intermediate_text.config(state=NORMAL)
            self.intermediate_text.delete("1.0", END)
            self.intermediate_text.insert(END, self._intermediate_results[stage_id])
            self.intermediate_text.config(state=DISABLED)

    def _update_intermediate_combo(self):
        """Обновляет список доступных этапов в выпадающем списке"""
        if hasattr(self, '_intermediate_results'):
            stages = list(self._intermediate_results.keys())
            current = self.intermediate_combo.get()
            self.intermediate_combo['values'] = stages
            # Если текущее значение есть в новом списке — оставляем его
            if current in stages:
                self.intermediate_combo.set(current)
            elif stages:
                # Иначе выбираем последний добавленный
                self.intermediate_combo.set(stages[-1])
                self._on_intermediate_selected(None)

    def load_models_at_startup(self):
        """Асинхронно загружает список моделей при запуске"""

        def load():
            try:
                api_key = self.config_manager.get_api_key()
                if not api_key:
                    self.logger.warning("Не удалось загрузить модели: нет API ключа")
                    return
                from openai import OpenAI
                client = OpenAI(api_key=api_key)
                models = client.models.list()
                # Фильтруем только GPT модели
                model_names = [m.id for m in models if 'gpt' in m.id and not m.id.startswith('whisper')]
                # Сортировка с приоритетом
                priority = ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo", "gpt-4"]
                sorted_models = sorted(model_names, key=lambda x: (
                    x not in priority, priority.index(x) if x in priority else len(priority)))
                self.available_models = sorted_models
                self.logger.info(f"Загружено {len(sorted_models)} моделей")
            except Exception as e:
                self.logger.error(f"Ошибка загрузки моделей: {e}")
                self.available_models = ["gpt-4o-mini", "gpt-4o", "gpt-4-turbo", "gpt-4"]  # fallback

        threading.Thread(target=load, daemon=True).start()

    def new_dialog(self):
        if self.dialog_manager.current_dialog["messages"]:
            reply = messagebox.askyesno("Новый диалог", "Текущий диалог будет потерян. Продолжить?")
            if not reply:
                return
        self.dialog_manager.new_dialog()
        self.chat_widget.clear()
        self.attached_files = []
        self._update_file_label()
        self.statusbar.config(text="Новый диалог создан")

    def save_dialog(self):
        if not self.dialog_manager.current_dialog["messages"]:
            messagebox.showwarning("Сохранение", "Нет сообщений для сохранения.")
            return
        # Добавляем текущие прикреплённые файлы в диалог
        for fname, content in self.attached_files:
            self.dialog_manager.add_attached_file(fname, content)
        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json")],
            initialdir=self.dialog_manager.dialogs_dir,
            initialfile=f"{self.dialog_manager.current_dialog['dialog_id']}.json"
        )
        if file_path:
            self.dialog_manager.save_dialog(file_path)
            self.statusbar.config(text=f"Диалог сохранён в {file_path}")

    def load_dialog(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("JSON files", "*.json")],
            initialdir=self.dialog_manager.dialogs_dir
        )
        if file_path:
            self.dialog_manager.load_dialog(file_path)
            # Очищаем GUI чат
            self.chat_widget.clear()
            # Восстанавливаем сообщения
            for msg in self.dialog_manager.current_dialog["messages"]:
                self.chat_widget.add_message(msg["content"], msg["role"])
            # Восстанавливаем прикреплённые файлы
            self.attached_files = [(f["filename"], f["content"]) for f in
                                   self.dialog_manager.current_dialog.get("attached_files", [])]
            self._update_file_label()
            self.statusbar.config(text=f"Диалог загружен: {file_path}")

    def summarize_dialog(self):
        if not self.dialog_manager.current_dialog["messages"]:
            messagebox.showwarning("Обобщение", "Нет сообщений для обобщения.")
            return
        self.statusbar.config(text="Генерация summary...")
        # Получаем последние N пар для summary (из настроек)
        dialog_params = self.config_manager.get_dialog_params()
        summary_pairs = dialog_params.get("summary_pairs", 3)
        pairs = self.dialog_manager.get_conversation_pairs(summary_pairs)
        # Формируем текст для summary: последние пары + предыдущий summary
        conv_text = "Предыдущее обобщение:\n" + self.dialog_manager.current_dialog.get("summary", "") + "\n\n"
        for i, (q, a) in enumerate(pairs):
            conv_text += f"Вопрос {i + 1}: {q}\nОтвет {i + 1}: {a}\n"
        # Добавляем содержимое файлов, загруженных в текущей сессии
        if self.attached_files:
            conv_text += "\nСодержимое загруженных файлов:\n"
            for fname, content in self.attached_files:
                conv_text += f"--- {fname} ---\n{content}\n"

        # Вызываем метод генерации summary в отдельном потоке
        def generate():
            summary = self.chat_client.generate_summary(conv_text)
            self.after(0, lambda: self._on_summary_ready(summary))

        threading.Thread(target=generate, daemon=True).start()

    def _on_summary_ready(self, summary):
        self.dialog_manager.current_dialog["summary"] = summary
        self.dialog_manager.save_dialog()  # автосохранение
        self.chat_widget.add_message(f"*Система*: Обобщение диалога:\n{summary}", "assistant")
        self.statusbar.config(text="Summary готов")

    # ---------- Обработчики колбэков (потокобезопасные) ----------
    def _on_ai_chunk(self, chunk):
        self.after(0, lambda: self._append_chunk(chunk))

    def _append_chunk(self, chunk):
        self._stream_buffer += chunk

    def _on_ai_response(self, text, metadata):
        self.after(0, lambda: self._display_response(text))

    def _display_response(self, text):
        full = text if text else self._stream_buffer
        if full:
            self.chat_widget.add_message(full, "assistant")
            self.dialog_manager.add_message("assistant", full, tokens_used=0)  # токены можно не сохранять
        else:
            self.chat_widget.add_message("*Ассистент не дал ответа*", "assistant")
        self._stream_buffer = ""

    def _on_api_error(self, msg, recoverable):
        self.after(0, lambda: self._show_error(msg, recoverable))

    def _show_error(self, msg, recoverable):
        self.logger.error(f"API error: {msg}")
        self.chat_widget.add_message(f"*Ошибка*: {msg}", "assistant")
        if not recoverable:
            self.statusbar.config(text="Критическая ошибка API. Проверьте ключ.")
        else:
            self.statusbar.config(text="Ошибка API, повторите попытку.")

    def _on_request_start(self):
        self.after(0, lambda: self._set_sending_state(True))

    def _on_request_finish(self):
        self.after(0, lambda: self._set_sending_state(False))

    def _set_sending_state(self, is_sending):
        if is_sending:
            self.chat_widget.set_send_enabled(False)
            self.statusbar.config(text="Ассистент печатает...")
        else:
            self.chat_widget.set_send_enabled(True)
            self.statusbar.config(text="Готов")

    # ---------- Пользовательские действия ----------
    def on_user_message(self, text: str):
        """Обрабатывает сообщение пользователя: команда или обычный вопрос."""
        if text.startswith('/'):
            self._handle_command(text)
            return

        self.chat_widget.add_message(text, "user")
        self.dialog_manager.add_message("user", text)

        # Получаем настройки диалога
        dialog_params = self.config_manager.get_dialog_params()
        history_pairs = dialog_params.get("history_pairs", 3)

        # Последние N пар
        pairs = self.dialog_manager.get_conversation_pairs(history_pairs)
        history_text = ""
        for i, (q, a) in enumerate(pairs):
            history_text += f"Вопрос {i + 1}: {q}\nОтвет {i + 1}: {a}\n"

        summary = self.dialog_manager.current_dialog.get("summary", "")

        # Получаем содержимое файлов (прикреплённые через кнопку или loaded_documents)
        files_text = ""
        loaded_docs = self.dialog_manager.get_loaded_documents()
        for doc in loaded_docs:
            files_text += f"--- {doc['filename']} ---\n{doc['content']}\n"

        # Также добавляем прикреплённые через кнопку (для обратной совместимости)
        for fname, content in self.attached_files:
            if not any(doc["filename"] == fname for doc in loaded_docs):
                files_text += f"--- {fname} ---\n{content}\n"

        # Загружаем промпты
        prompts = self.chat_client._load_prompts()
        system_prompt = prompts.get("dialog_system_prompt", "Ты полезный ассистент.")
        user_prompt_template = prompts.get(
            "dialog_user_prompt",
            "Краткое содержание предыдущих сообщений (summary):\n{summary}\n\n"
            "Последние вопросы и ответы:\n{history}\n\n"
            "Содержимое загруженных файлов:\n{files}\n\n"
            "Текущий вопрос пользователя: {question}"
        )
        user_prompt = user_prompt_template.format(
            summary=summary,
            history=history_text,
            files=files_text,
            question=text
        )

        # Получаем настройки API
        api_key = self.config_manager.get_api_key()
        if not api_key:
            self.chat_widget.add_message("❌ API ключ не найден. Задайте его в настройках.", "assistant")
            return

        llm_params = self.config_manager.get_llm_params()
        api_settings = {
            "api_key": api_key,
            "model": llm_params.get("model", "gpt-4o-mini"),
            "temperature": llm_params.get("temperature", 0.2),
            "max_tokens": llm_params.get("max_tokens", 2000),
            "timeout": llm_params.get("timeout", 60)
        }

        # Проверка размера файла
        if files_text:
            import tiktoken
            try:
                llm_params = self.config_manager.get_llm_params()
                model = llm_params.get("model", "gpt-4o-mini")
                encoding = tiktoken.encoding_for_model(model)
                total_tokens = len(encoding.encode(files_text))

                self.logger.info(f"📊 Размер файла в токенах: {total_tokens}")

                # Определяем безопасный лимит
                if "gpt-4" in model or "gpt-5" in model:
                    max_safe_tokens = 100000
                else:
                    max_safe_tokens = 50000

                if total_tokens > max_safe_tokens:
                    recommendation = self._recommend_file_split(files_text, model)
                    if recommendation:
                        self.chat_widget.add_message(recommendation, "assistant")
                        self.dialog_manager.add_message("assistant", recommendation)
                    else:
                        self.chat_widget.add_message(
                            "⚠️ Файл слишком большой. Попробуйте разбить его на части и загружать по отдельности.",
                            "assistant"
                        )
                    return  # Не отправляем запрос
            except Exception as e:
                self.logger.warning(f"Ошибка при оценке размера файла: {e}. Продолжаем обычную обработку.")

        # Обычный режим (файл небольшой или нет файлов)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]

        self._stream_buffer = ""
        self.chat_client.send_message(messages, stream=True)

    def _handle_command(self, text: str):
        """Обрабатывает команды, начинающиеся с '/'."""
        parts = text.split(maxsplit=1)
        cmd = parts[0].lower()
        args = parts[1] if len(parts) > 1 else ""

        if cmd == "/help":
            self._cmd_help()
        elif cmd == "/load_document":
            self._cmd_load_document(args)
        elif cmd == "/load_scenario":
            self._cmd_load_scenario(args)
        elif cmd == "/run_scenario":
            self._cmd_run_scenario()
        elif cmd == "/export_result":
            self._cmd_export_result(args)
        elif cmd == "/import_folder":
            self._cmd_import_folder(args)
        elif cmd == "/clear_documents":
            self._cmd_clear_documents()
        elif cmd == "/clear_scenarios":
            self._cmd_clear_scenarios()
        else:
            self.chat_widget.add_message(f"❌ Неизвестная команда: {cmd}. Введите /help для списка команд.", "assistant")
            self.dialog_manager.add_message("assistant", f"Неизвестная команда: {cmd}")

    def _cmd_load_document(self, file_path: str):
        """Загружает один или несколько документов."""
        if not file_path:
            self.chat_widget.add_message("❌ Укажите путь к файлу: /load_document <путь>", "assistant")
            return

        # Проверяем, является ли путь папкой или файлом
        if os.path.isdir(file_path):
            self._cmd_import_folder(file_path)
            return

        # Если несколько файлов через запятую или пробел
        if ',' in file_path or (' ' in file_path and '*' not in file_path):
            paths = [p.strip().strip('"').strip("'") for p in file_path.replace(',', ' ').split() if p.strip()]
            count = 0
            for p in paths:
                if os.path.exists(p):
                    self._load_single_document(p)
                    count += 1
                else:
                    self.chat_widget.add_message(f"⚠️ Файл не найден: {p}", "assistant")
            if count > 0:
                self.chat_widget.add_message(f"✅ Загружено {count} документов", "assistant")
            return

        # Одиночный файл
        self._load_single_document(file_path)

    def _cmd_load_document_dialog(self):
        """Открывает диалог выбора документа и загружает его."""
        file_path = filedialog.askopenfilename(
            title="Выберите документ для загрузки",
            filetypes=[
                ("Поддерживаемые", "*.docx *.xlsx *.pdf *.xls *.txt *.md *.json *.csv"),
                ("Документы Word", "*.docx"),
                ("Таблицы Excel", "*.xlsx *.xls"),
                ("PDF", "*.pdf"),
                ("Текстовые", "*.txt *.md *.json *.csv"),
                ("Все файлы", "*.*")
            ]
        )
        if file_path:
            self._load_single_document(file_path)

    def _load_single_document(self, file_path: str):
        """Загружает один документ."""
        if not os.path.exists(file_path):
            self.chat_widget.add_message(f"⚠️ Файл не найден: {file_path}", "assistant")
            return

        try:
            result = FileLoader._read_file(file_path, self.config_manager)
            if result:
                filename, content, file_format = result
                self.dialog_manager.add_loaded_document(filename, content, file_format, file_path)
                self.attached_files.append((filename, content))
                self._update_file_label()
                self.chat_widget.add_message(f"✅ Документ '{filename}' загружен в контекст", "assistant")
                self.dialog_manager.add_message("assistant", f"Документ '{filename}' загружен в контекст")
            else:
                self.chat_widget.add_message(f"❌ Не удалось загрузить документ: {file_path}", "assistant")
        except Exception as e:
            self.chat_widget.add_message(f"❌ Ошибка загрузки: {e}", "assistant")

    def _cmd_load_scenario(self, file_path: str):
        """Загружает сценарий по указанному пути."""
        if not file_path:
            self.chat_widget.add_message("❌ Укажите путь к сценарию: /load_scenario <путь>", "assistant")
            return

        file_path = file_path.strip().strip('"').strip("'")
        if not os.path.exists(file_path):
            self.chat_widget.add_message(f"❌ Файл сценария не найден: {file_path}", "assistant")
            return

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if 'stages' not in data:
                    raise ValueError("Нет ключа 'stages'")
            self.dialog_manager.add_loaded_scenario(file_path, data)
            self.chat_widget.add_message(f"✅ Сценарий '{Path(file_path).stem}' загружен", "assistant")
            self.dialog_manager.add_message("assistant", f"Сценарий '{Path(file_path).stem}' загружен")
        except Exception as e:
            self.chat_widget.add_message(f"❌ Ошибка загрузки сценария: {e}", "assistant")

    def _cmd_run_scenario(self):
        """Выполняет загруженный сценарий на загруженном документе."""
        scenario = self.dialog_manager.get_last_scenario()
        document = self.dialog_manager.get_last_document()

        if not scenario:
            self.chat_widget.add_message("❌ Сценарий не загружен. Используйте /load_scenario", "assistant")
            return

        if not document:
            self.chat_widget.add_message("❌ Документ не загружен. Используйте /load_document", "assistant")
            return

        # Получаем полный путь к документу
        doc_path = document.get("path")
        if not doc_path or not os.path.exists(doc_path):
            # Если путь не сохранён или файл удалён, пробуем найти по имени
            self.chat_widget.add_message(f"⚠️ Путь к документу не найден. Используйте /load_document с полным путём",
                                         "assistant")
            return

        # Получаем путь к сценарию
        scenario_path = scenario.get("path")
        if not scenario_path or not os.path.exists(scenario_path):
            self.chat_widget.add_message(f"⚠️ Путь к сценарию не найден. Используйте /load_scenario с полным путём",
                                         "assistant")
            return

        # Получаем настройки API
        api_key = self.config_manager.get_api_key()
        if not api_key:
            self.chat_widget.add_message("❌ API ключ не найден. Задайте его в настройках.", "assistant")
            return

        llm_params = self.config_manager.get_llm_params()
        api_settings = {
            "api_key": api_key,
            "model": llm_params.get("model", "gpt-4o-mini"),
            "temperature": llm_params.get("temperature", 0.2),
            "max_tokens": llm_params.get("max_tokens", 2000),
            "timeout": llm_params.get("timeout", 60)
        }

        self.chat_widget.add_message("🔄 Запуск анализа по сценарию...", "assistant")
        self.dialog_manager.add_message("assistant", "Запуск анализа по сценарию...")

        # Запускаем в отдельном потоке
        self.cancel_event = threading.Event()
        self.analysis_thread = threading.Thread(
            target=self._run_scenario_thread,
            args=(doc_path, scenario_path, api_settings),
            daemon=True
        )
        self.analysis_thread.start()

    def _run_scenario_thread(self, doc_path, scenario_path, api_settings):
        """Выполняет сценарий в отдельном потоке с выводом в чат."""

        def chat_callback(msg_type, data):
            # Преобразуем сообщения для вывода в чат
            if msg_type == "log":
                self.after(0, lambda: self._append_chat_message(f"*Система*: {data}"))
            elif msg_type == "stage":
                current = data.get("current", 0)
                total = data.get("total", 1)
                name = data.get("name", "unknown")
                self.after(0, lambda: self._append_chat_message(f"*Система*: Этап {current}/{total}: {name}"))
            elif msg_type == "stage_result":
                stage_id = data.get("stage_id", "unknown")
                preview = data.get("result_preview", "")
                self.after(0, lambda: self._append_chat_message(f"*Система*: Результат этапа '{stage_id}':\n{preview}"))

        try:
            result = run_analysis(
                document_path=doc_path,
                scenario_path=scenario_path,
                api_settings=api_settings,
                progress_callback=chat_callback,
                cancel_event=self.cancel_event
            )
            self.after(0, lambda: self._on_scenario_finished(result))
        except Exception as e:
            self.after(0, lambda: self._append_chat_message(f"❌ Критическая ошибка: {e}"))

    def _append_chat_message(self, text: str):
        """Добавляет сообщение в чат от имени системы."""
        self.chat_widget.add_message(text, "assistant")
        self.dialog_manager.add_message("assistant", text)

    def _on_scenario_finished(self, result):
        """Обрабатывает завершение выполнения сценария."""
        if result["status"] == "success":
            self.dialog_manager.set_last_result(result["result"], result.get("format", "text"))
            self._append_chat_message(f"✅ Анализ завершён успешно! Результат в формате {result.get('format', 'text')}")
            # Показываем результат в чате (если не слишком большой)
            preview = result["result"][:1000] + ("..." if len(result["result"]) > 1000 else "")
            self._append_chat_message(f"**Результат:**\n{preview}")
        elif result["status"] == "cancelled":
            self._append_chat_message("⚠️ Анализ отменён пользователем")
        else:
            self._append_chat_message(f"❌ Ошибка: {result.get('message', 'Неизвестная ошибка')}")

        self.analysis_thread = None

    def _cmd_export_result(self, format_type: str):
        """Экспортирует последний результат анализа в указанном формате."""
        if not format_type:
            self.chat_widget.add_message("❌ Укажите формат: /export_result <json|csv|md|docx|xlsx>", "assistant")
            return

        format_type = format_type.strip().lower()
        valid_formats = {"json", "csv", "md", "docx", "xlsx"}
        if format_type not in valid_formats:
            self.chat_widget.add_message(
                f"❌ Неподдерживаемый формат: {format_type}. Доступные: {', '.join(valid_formats)}", "assistant")
            return

        result = self.dialog_manager.get_last_result()
        if not result:
            self.chat_widget.add_message("❌ Нет сохранённых результатов анализа", "assistant")
            return

        # Маппинг форматов к расширениям
        ext_map = {
            "json": ".json",
            "csv": ".csv",
            "md": ".md",
            "docx": ".docx",
            "xlsx": ".xlsx"
        }

        file_path = filedialog.asksaveasfilename(
            defaultextension=ext_map[format_type],
            filetypes=[(f"{format_type.upper()} files", f"*{ext_map[format_type]}"), ("All files", "*.*")],
            title=f"Сохранить результат как {format_type.upper()}"
        )
        if not file_path:
            return

        try:
            ext = Path(file_path).suffix.lower()
            if ext == '.docx':
                self._save_as_docx(file_path, result["result"], result.get("format", "text"))
            elif ext == '.xlsx':
                self._save_as_xlsx(file_path, result["result"], result.get("format", "text"))
            else:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(result["result"])
            self.chat_widget.add_message(f"✅ Результат сохранён: {file_path}", "assistant")
            self.dialog_manager.add_message("assistant", f"Результат сохранён: {file_path}")
            self.statusbar.config(text=f"Результат сохранён: {file_path}")
        except Exception as e:
            self.chat_widget.add_message(f"❌ Ошибка сохранения: {e}", "assistant")
            self.logger.error(f"Ошибка экспорта: {e}")

    def _cmd_import_folder(self, folder_path: str):
        """Загружает все поддерживаемые документы из папки."""
        if not folder_path:
            self.chat_widget.add_message("❌ Укажите путь к папке: /import_folder <путь>", "assistant")
            return

        folder_path = folder_path.strip().strip('"').strip("'")
        if not os.path.isdir(folder_path):
            self.chat_widget.add_message(f"❌ Папка не найдена: {folder_path}", "assistant")
            return

        supported_ext = {'.docx', '.xlsx', '.pdf', '.xls', '.txt', '.md', '.json', '.csv'}
        count = 0
        errors = []

        for ext in supported_ext:
            for file_path in glob.glob(os.path.join(folder_path, f"*{ext}")):
                try:
                    result = FileLoader._read_file(file_path, self.config_manager)
                    if result:
                        filename, content, file_format = result
                        self.dialog_manager.add_loaded_document(filename, content, file_format,
                                                                file_path)  # передаём путь
                        self.attached_files.append((filename, content))
                        count += 1
                except Exception as e:
                    errors.append(f"{Path(file_path).name}: {e}")

        self._update_file_label()
        if count > 0:
            self.chat_widget.add_message(f"✅ Загружено {count} документов из папки", "assistant")
            self.dialog_manager.add_message("assistant", f"Загружено {count} документов из папки")
        if errors:
            self.chat_widget.add_message(f"⚠️ Ошибки при загрузке:\n" + "\n".join(errors), "assistant")

    def _cmd_clear_documents(self):
        """Очищает загруженные документы."""
        self.dialog_manager.clear_loaded_documents()
        self.attached_files = []
        self._update_file_label()
        self.chat_widget.add_message("✅ Загруженные документы очищены", "assistant")
        self.dialog_manager.add_message("assistant", "Загруженные документы очищены")

    def _cmd_clear_scenarios(self):
        """Очищает загруженные сценарии."""
        self.dialog_manager.clear_loaded_scenarios()
        self.chat_widget.add_message("✅ Загруженные сценарии очищены", "assistant")
        self.dialog_manager.add_message("assistant", "Загруженные сценарии очищены")

    def _cmd_load_document(self, file_path: str):
        """Загружает один или несколько документов (поддерживает пути с пробелами)."""
        if not file_path:
            self.chat_widget.add_message("❌ Укажите путь к файлу: /load_document <путь>", "assistant")
            return

        file_path = file_path.strip()

        # Проверяем, является ли путь папкой
        if os.path.isdir(file_path):
            self._cmd_import_folder(file_path)
            return

        # Обработка множественных путей (разделитель — запятая или точка с запятой)
        # Это позволяет использовать пути с пробелами внутри
        if ',' in file_path or ';' in file_path:
            separator = ',' if ',' in file_path else ';'
            paths = [p.strip().strip('"').strip("'") for p in file_path.split(separator) if p.strip()]
            count = 0
            for p in paths:
                # Удаляем кавычки, если есть
                p = p.strip('"').strip("'")
                if os.path.exists(p):
                    self._load_single_document(p)
                    count += 1
                else:
                    # Проверяем, может быть это имя файла без пути
                    self.chat_widget.add_message(f"⚠️ Файл не найден: {p}", "assistant")
            if count > 0:
                self.chat_widget.add_message(f"✅ Загружено {count} документов", "assistant")
            return

        # Одиночный файл (убираем кавычки, если есть)
        file_path = file_path.strip('"').strip("'")
        self._load_single_document(file_path)

    def _cmd_load_scenario_dialog(self):
        """Открывает диалог выбора сценария и загружает его."""
        file_path = filedialog.askopenfilename(
            title="Выберите JSON-сценарий",
            filetypes=[("JSON files", "*.json")]
        )
        if file_path:
            self._cmd_load_scenario(file_path)

    def _cmd_import_folder_dialog(self):
        """Открывает диалог выбора папки и импортирует все документы."""
        folder_path = filedialog.askdirectory(
            title="Выберите папку с документами"
        )
        if folder_path:
            self._cmd_import_folder(folder_path)

        # Создаём диалог выбора формата
        dialog = tk.Toplevel(self)
        dialog.title("Экспорт результата")
        dialog.geometry("300x200")
        dialog.transient(self)
        dialog.grab_set()

        tk.Label(dialog, text="Выберите формат:").pack(pady=10)

        format_var = tk.StringVar(value="json")
        formats = [("JSON", "json"), ("CSV", "csv"), ("Markdown", "md"),
                   ("Word DOCX", "docx"), ("Excel XLSX", "xlsx")]

        for text, value in formats:
            tk.Radiobutton(dialog, text=text, variable=format_var, value=value).pack(anchor=tk.W, padx=20)

        def on_export():
            dialog.destroy()
            self._cmd_export_result(format_var.get())

        btn_frame = tk.Frame(dialog)
        btn_frame.pack(pady=20)
        tk.Button(btn_frame, text="Экспортировать", command=on_export).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="Отмена", command=dialog.destroy).pack(side=tk.LEFT, padx=5)

    def _cmd_help(self):
        """Показывает справку по командам."""
        help_text = (
            "**Доступные команды:**\n\n"
            "`/load_document <путь>` — загрузить документ в контекст чата\n"
            "`/load_scenario <путь>` — загрузить JSON-сценарий\n"
            "`/run_scenario` — выполнить загруженный сценарий на загруженном документе\n"
            "`/export_result <формат>` — экспортировать результат (json, csv, md, docx, xlsx)\n"
            "`/import_folder <путь>` — загрузить все документы из папки\n"
            "`/clear_documents` — очистить загруженные документы\n"
            "`/clear_scenarios` — очистить загруженные сценарии\n"
            "`/help` — показать эту справку"
        )
        self.chat_widget.add_message(help_text, "assistant")
        self.dialog_manager.add_message("assistant", help_text)

    def _load_prompt_global(self):
        self.chat_widget._load_prompt_from_dialog()

    def _save_prompt_global(self):
        self.chat_widget._save_prompt_to_dialog()

    def _auto_load_prompt(self):
        prompt_defaults = self.config_manager.get_prompt_defaults()
        auto_file = prompt_defaults.get("auto_load_file", "")
        if auto_file and os.path.exists(auto_file):
            try:
                with open(auto_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.chat_widget.input_field.delete("1.0", END)
                self.chat_widget.input_field.insert("1.0", content)
                self.chat_widget.add_message(f"*Система*: Автозагрузка промпта из {auto_file}", "assistant")
            except Exception as e:
                self.logger.error(f"Ошибка автозагрузки промпта: {e}")

    def load_file(self):
        """Загружает файл через диалог и добавляет его в контекст чата."""
        result = FileLoader.load_from_dialog(self, self.config_manager)
        if result:
            filename, content, file_format = result
            self.attached_files.append((filename, content))
            self.dialog_manager.add_attached_file(filename, content)

            # Пытаемся найти полный путь к файлу
            file_path = self._find_file_by_name(filename)
            self.dialog_manager.add_loaded_document(filename, content, file_format, file_path)

            self._update_file_label()
            self.statusbar.config(text=f"Файл '{filename}' добавлен")
            self.chat_widget.add_message(f"*Система*: Загружен файл \"{filename}\".", "assistant")

            # Сохраняем извлечённое содержимое на диск
            saved_path = self._save_converted_file(filename, content, file_format)
            if saved_path:
                self.chat_widget.add_message(f"*Система*: Извлечённый текст сохранён в {saved_path}", "assistant")
        else:
            messagebox.showerror("Ошибка",
                                 "Не удалось загрузить файл.\nПоддерживаются: docx, xlsx, pdf, txt, md, py, json, csv")

    def _find_file_by_name(self, filename: str) -> str:
        """
        Пытается найти файл по имени в текущей директории и в папке documents.
        Возвращает полный путь или имя файла, если не найден.
        """
        # Проверяем текущую директорию
        if os.path.exists(filename):
            return filename

        # Проверяем папку documents
        docs_dir = self.config_manager.config.get("input_settings", {}).get("default_input_directory", "./documents")
        full_path = os.path.join(docs_dir, filename)
        if os.path.exists(full_path):
            return full_path

        # Проверяем папку converted_files (куда сохраняются конвертированные файлы)
        converted_dir = self.config_manager.get_paths().get("converted_files_dir", "./converted_files")
        full_path = os.path.join(converted_dir, filename)
        if os.path.exists(full_path):
            return full_path

        # Если не найден, возвращаем имя файла (путь будет невалидным, но это лучше чем ничего)
        return filename

    def _update_file_label(self):
        if self.attached_files:
            names = ", ".join(f[0] for f in self.attached_files)
            self.file_label.config(text=f"📎 Файлы: {names}")
            self.clear_file_btn.config(state=NORMAL)
        else:
            self.file_label.config(text="Файлы не загружены")
            self.clear_file_btn.config(state=DISABLED)

    def _save_converted_file(self, original_filename: str, content: str, file_format: str) -> Optional[str]:
        """
        Сохраняет извлечённое содержимое в папку converted_files.
        Возвращает путь к сохранённому файлу или None.
        """
        try:
            save_dir = self.config_manager.get_paths().get("converted_files_dir", "./converted_files")
            Path(save_dir).mkdir(parents=True, exist_ok=True)

            base = Path(original_filename).stem
            if file_format == "json":
                ext = ".json"
            else:
                ext = ".md"  # для markdown и text

            save_path = Path(save_dir) / f"{base}{ext}"
            counter = 1
            while save_path.exists():
                save_path = Path(save_dir) / f"{base}_{counter}{ext}"
                counter += 1

            with open(save_path, 'w', encoding='utf-8') as f:
                f.write(content)

            self.logger.info(f"Сохранён конвертированный файл: {save_path}")
            return str(save_path)
        except Exception as e:
            self.logger.error(f"Ошибка сохранения конвертированного файла: {e}")
            return None

    def clear_file(self):
        self.attached_files = []
        self.dialog_manager.clear_attached_files()
        self._update_file_label()
        self.statusbar.config(text="Все файлы удалены")
        self.chat_widget.add_message("*Система*: Все загруженные файлы удалены.", "assistant")

    def copy_chat(self):
        chat_text = self.chat_widget.get_chat_text()
        if chat_text:
            self.clipboard_clear()
            self.clipboard_append(chat_text)
            self.statusbar.config(text="Чат скопирован в буфер обмена")
            self.chat_widget.add_message("*Система*: Содержимое чата скопировано в буфер обмена.", "assistant")
        else:
            self.statusbar.config(text="Чат пуст, нечего копировать")

    def _save_chat_as_docx_wrapper(self):
        """Получает текст чата и сохраняет в DOCX."""
        chat_text = self.chat_widget.get_chat_text()
        if not chat_text:
            messagebox.showwarning("Сохранение", "Чат пуст, нечего сохранять")
            return
        self._save_chat_as_docx(chat_text)

    def _save_chat_as_docx(self, chat_text: str):
        """Сохраняет содержимое чата в DOCX с поддержкой Markdown-разметки."""
        doc = Document()

        # Парсим текст чата построчно
        lines = chat_text.split('\n')
        in_code_block = False
        code_block_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            # Пропускаем разделители (---)
            if line.startswith('---') and line.endswith('---'):
                continue

            # Обработка маркеров чата (Вы, Ассистент, Система)
            if line.startswith('Вы  ') or line.startswith('Ассистент  ') or line.startswith('*Система*'):
                # Это заголовок сообщения, делаем жирным
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(10)
                continue

            # Обработка кодовых блоков
            if line.startswith('```'):
                if in_code_block:
                    # Закрываем блок
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_block_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    code_block_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_block_lines.append(line)
                continue

            # Обработка Markdown-элементов
            p = doc.add_paragraph()

            # Заголовки
            if line.startswith('# '):
                run = p.add_run(line[2:])
                run.bold = True
                run.font.size = Pt(18)
                continue
            elif line.startswith('## '):
                run = p.add_run(line[3:])
                run.bold = True
                run.font.size = Pt(16)
                continue
            elif line.startswith('### '):
                run = p.add_run(line[4:])
                run.bold = True
                run.font.size = Pt(14)
                continue
            elif line.startswith('#### '):
                run = p.add_run(line[5:])
                run.bold = True
                run.font.size = Pt(12)
                continue

            # Обработка Markdown: жирный (**текст**), курсив (*текст*), код (`текст`)
            # Жирный
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Курсив
                    subparts = re.split(r'(\*[^*]+\*)', part)
                    for subpart in subparts:
                        if subpart.startswith('*') and subpart.endswith('*'):
                            run = p.add_run(subpart[1:-1])
                            run.italic = True
                        else:
                            # Моноширинный код (`текст`)
                            code_parts = re.split(r'(`[^`]+`)', subpart)
                            for code_part in code_parts:
                                if code_part.startswith('`') and code_part.endswith('`'):
                                    run = p.add_run(code_part[1:-1])
                                    run.font.name = 'Courier New'
                                else:
                                    p.add_run(code_part)

        # Сохраняем файл
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Документ Word", "*.docx"), ("Все файлы", "*.*")],
            title="Сохранить чат как DOCX"
        )
        if file_path:
            doc.save(file_path)
            self.statusbar.config(text=f"Чат сохранён в DOCX: {file_path}")

    def check_api_key(self):
        if not self.config_manager.get_api_key():
            reply = messagebox.askyesno("API ключ", "OPENAI_API_KEY не найден. Ввести сейчас?")
            if reply:
                from tkinter import simpledialog
                key = simpledialog.askstring("API ключ", "Введите OpenAI API ключ:", show='*')
                if key:
                    self.config_manager.set_api_key(key)
                    self.statusbar.config(text="API ключ сохранён")
                    self.chat_widget.add_message("*Система*: API ключ успешно сохранён.", "assistant")
                else:
                    self.statusbar.config(text="API ключ не задан")
            else:
                self.statusbar.config(text="API ключ не задан")

    def test_api_connection(self):
        def test():
            success = self.chat_client.test_connection()
            self.after(0, lambda: self._on_test_result(success))

        threading.Thread(target=test, daemon=True).start()

    def _on_test_result(self, success):
        if success:
            self.chat_widget.add_message("*Система*: ✅ Соединение с OpenAI API установлено.", "assistant")
            self.statusbar.config(text="API готов")
        else:
            self.chat_widget.add_message("*Система*: ❌ Не удалось подключиться к API. Проверьте ключ и интернет.",
                                         "assistant")
            self.statusbar.config(text="Ошибка подключения к API")

    def clear_chat(self):
        self.chat_widget.clear()
        self.statusbar.config(text="Чат очищен")
        self.chat_widget.add_message("*Система*: История чата очищена.", "assistant")

    def open_settings(self):
        dialog = SettingsDialog(self, self.config_manager, self.available_models)
        self.wait_window(dialog)
        theme = self.config_manager.config.get("ui", {}).get("theme", "light")
        new_theme = "darkly" if theme == "dark" else "cosmo"
        self.style.theme_use(new_theme)

    def show_about(self):
        messagebox.showinfo(
            "О программе",
            "AI Document Analyst (Chat)\nВерсия 1.0\n\n"
            "Агентная система анализа документов с ИИ.\n"
            "Использует OpenAI API для ответов на вопросы.\n\n"
            "© 2024"
        )


    def _recommend_file_split(self, content: str, model: str = "gpt-4o-mini") -> str:
        """Подсчитывает токены и возвращает рекомендацию по разбиению файла."""
        try:
            import tiktoken
            encoding = tiktoken.encoding_for_model(model)
            total_tokens = len(encoding.encode(content))

            # Безопасный лимит для модели (учитываем промпт и ответ)
            if "gpt-4" in model or "gpt-5" in model:
                max_safe_tokens = 100000  # оставляем запас
            else:
                max_safe_tokens = 50000  # для gpt-3.5

            if total_tokens <= max_safe_tokens:
                return None  # файл помещается

            # Рекомендуемое количество частей
            parts = (total_tokens // max_safe_tokens) + 1
            tokens_per_part = total_tokens // parts + 1

            return (
                f"⚠️ **Файл слишком большой для модели {model}.**\n"
                f"Размер: ~{total_tokens:,} токенов.\n"
                f"Рекомендация: разбейте файл на **{parts}** частей (примерно по {tokens_per_part:,} токенов каждая)\n"
                f"и загружайте их по очереди, задавая вопросы к каждой части отдельно.\n"
                f"Или используйте команду /load_document для загрузки отдельных файлов."
            )
        except Exception as e:
            self.logger.error(f"Ошибка подсчёта токенов: {e}")
            return None
