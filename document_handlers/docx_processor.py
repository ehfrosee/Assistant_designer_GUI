# [file name]: docx_processor.py
import os
import re
from pathlib import Path
from typing import List, Dict, Any

from document_handlers.base_processor import BaseDocumentProcessor

try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Paragraph = None
    Table = None


class DOCXProcessor(BaseDocumentProcessor):
    """Процессор для DOCX файлов с поддержкой колонтитулов"""

    def __init__(self, config: Dict[str, Any]):
        super().__init__(config)
        self.processor_config = config.get('processors', {}).get('docx', {})
        self.extract_headers_footers = self.processor_config.get('extract_headers_footers', True)
        self.extract_tables = self.processor_config.get('extract_tables', True)
    
    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith('.docx') and HAS_DOCX
    
    def extract_content(self, file_path: str) -> str:
        """Извлечение содержимого из DOCX"""
        doc = Document(file_path)
        content_parts = []
        
        # Извлечение колонтитулов
        if self.extract_headers_footers:
            header_text = self._extract_headers(doc)
            if header_text:
                content_parts.append("### Верхние колонтитулы")
                content_parts.append(header_text)
                content_parts.append("")
            
            footer_text = self._extract_footers(doc)
            if footer_text:
                stamp_data = self._parse_stamp(footer_text)
                if stamp_data:
                    content_parts.append("### Штамп документа")
                    for key, value in stamp_data.items():
                        content_parts.append(f"- **{key}**: {value}")
                    content_parts.append("")
        
        # Извлечение основного содержимого
        for element in self._get_elements_ordered(doc):
            if element['type'] == 'paragraph':
                text = self._format_paragraph(element['object'])
                if text:
                    content_parts.append(text)
            
            elif element['type'] == 'table' and self.extract_tables:
                table_text = self._format_table(element['object'])
                if table_text:
                    content_parts.append(table_text)
        
        return '\n'.join(content_parts)
    
    def _extract_headers(self, doc) -> str:
        """Извлечение текста из верхних колонтитулов"""
        headers = []
        
        for section in doc.sections:
            if section.header:
                text = self._get_paragraphs_text(section.header.paragraphs)
                if text:
                    headers.append(text)
            
            if hasattr(section, 'first_page_header') and section.first_page_header:
                text = self._get_paragraphs_text(section.first_page_header.paragraphs)
                if text:
                    headers.append(text)
        
        return '\n'.join(headers)
    
    def _extract_footers(self, doc) -> str:
        """Извлечение текста из нижних колонтитулов"""
        footers = []
        
        for section in doc.sections:
            if section.footer:
                text = self._get_paragraphs_text(section.footer.paragraphs)
                if text:
                    footers.append(text)
        
        return '\n'.join(footers)
    
    def _get_paragraphs_text(self, paragraphs) -> str:
        """Получение текста из параграфов"""
        return '\n'.join(p.text.strip() for p in paragraphs if p.text.strip())
    
    def _parse_stamp(self, footer_text: str) -> Dict[str, str]:
        """Парсинг штампа из нижнего колонтитула"""
        stamp = {}
        
        # Поиск шифра проекта
        code_match = re.search(r'(\d{2}-\d{2}-[А-Яа-я0-9\.]+)', footer_text)
        if code_match:
            stamp['document_code'] = code_match.group(1)
        
        # Поиск даты
        date_match = re.search(r'(\d{2}\.\d{4}|\d{4}-\d{2}-\d{2})', footer_text)
        if date_match:
            stamp['date'] = date_match.group(1)
        
        # Поиск названия объекта
        obj_match = re.search(r'Объект[:\s]+(.+)$', footer_text, re.MULTILINE)
        if obj_match:
            stamp['project_name'] = obj_match.group(1).strip()
        
        # Поиск организации
        org_match = re.search(r'(ООО|АО|ЗАО|ПАО)\s+[\w\s"\'-]+', footer_text)
        if org_match:
            stamp['organization'] = org_match.group(0).strip()
        
        return stamp
    
    def _get_elements_ordered(self, doc) -> List[Dict]:
        """Получение элементов документа в правильном порядке"""
        elements = []
        
        try:
            for block in doc.element.body:
                if block.tag.endswith('p'):
                    paragraph = Paragraph(block, doc)
                    if paragraph.text.strip():
                        elements.append({'type': 'paragraph', 'object': paragraph})
                
                elif block.tag.endswith('tbl'):
                    table = Table(block, doc)
                    if self._has_table_content(table):
                        elements.append({'type': 'table', 'object': table})
        
        except Exception:
            # Fallback
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    elements.append({'type': 'paragraph', 'object': paragraph})
            
            for table in doc.tables:
                if self._has_table_content(table):
                    elements.append({'type': 'table', 'object': table})
        
        return elements
    
    def _has_table_content(self, table) -> bool:
        """Проверка наличия содержимого в таблице"""
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    return True
        return False
    
    def _format_paragraph(self, paragraph) -> str:
        """Форматирование параграфа"""
        text = paragraph.text.strip()
        if not text:
            return ""
        
        # Определение заголовков
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            
            if 'heading 1' in style_name or 'title' in style_name:
                return f"# {text}"
            elif 'heading 2' in style_name:
                return f"## {text}"
            elif 'heading 3' in style_name:
                return f"### {text}"
        
        # Проверка на жирный текст
        if self._is_all_bold(paragraph) and len(text) < 100:
            return f"### {text}"
        
        return text
    
    def _is_all_bold(self, paragraph) -> bool:
        """Проверка, является ли весь параграф жирным"""
        if not paragraph.runs:
            return False
        
        for run in paragraph.runs:
            if run.text.strip() and not run.bold:
                return False
        
        return any(run.text.strip() for run in paragraph.runs)
    
    def _format_table(self, table) -> str:
        """Форматирование таблицы"""
        table_data = []
        
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                table_data.append(row_data)
        
        return self.table_processor.format_table(table_data)